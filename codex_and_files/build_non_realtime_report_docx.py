from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "codex_and_files" / "non_realtime_experiment_summary_2026-08-24"
DATA = json.loads((OUT / "performance_summary.json").read_text(encoding="utf-8"))
DOCX = OUT / "Objective3_non_realtime_experiments_detailed_summary_2026-08-24.docx"


NAVY = "1F3A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GRID = "C9D2DC"
WHITE = "FFFFFF"
RISK = "9B1C1C"
GOLD = "7A5A00"

MODEL_NAMES = {
    "m0": "M0（当前 clip）",
    "m1": "M1（history，无位置）",
    "m2": "M2（实际顺序+位置）",
    "m3": "M3（固定 graph-valid+位置）",
    "m4": "M4（candidate，无 relation）",
    "m5": "M5（oracle relation）",
    "m6": "M6（soft relation）",
    "e2e_tier3_scratch": "E2E Tier3 Scratch",
    "e2e_node_scratch": "E2E Node Scratch",
    "e2e_node_from_tier3": "E2E Node From Tier3",
    "m1_direct": "M1 Direct",
    "m2_direct": "M2 Direct",
    "m3_direct": "M3 Direct",
    "m3_dynamic_frozen_m0_delta": "Dynamic Frozen-M0 Delta",
    "m3_dynamic_joint_head_delta": "Dynamic Joint-Head Delta",
    "m3_dynamic_direct_fusion": "Dynamic Direct Fusion",
    "refresh_every_1": "Atomic 每 1 epoch 刷新",
    "refresh_every_10": "Atomic 每 10 epoch 刷新",
    "refresh_once": "Atomic 只生成一次",
}


def set_run_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    field_run = paragraph.add_run()
    set_run_font(field_run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        field_run._r.append(node)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MID_GRAY)


def configure_document(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        r = hp.add_run("Objective 3 · Non-real-time Experiment Record")
        set_run_font(r, size=9, bold=True, color=MID_GRAY)
        add_page_field(sec.footer.paragraphs[0])


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("EXPERIMENT SUMMARY REPORT")
    set_run_font(r, size=10.5, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Objective 3：非 real-time 实验全流程总结")
    set_run_font(r, size=26, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Graph-history、Graph-valid Shuffle、Atomic-tail 与 DualPos 消融")
    set_run_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("覆盖：J 先导 → 严格四折 → Direct → Dynamic → Atomic-tail → A0–A4 DualPos")
    set_run_font(r, size=10.5, italic=True, color=MID_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("整理日期：2026-08-24  |  相机：001484412812  |  不含实时边界分割与在线识别实验")
    set_run_font(r, size=10, color=MID_GRAY)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text="", bold_prefix=None, color=None, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)

    for text in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_node])
        p_pr.insert(0, num_pr)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r)


def add_table(doc, headers, rows, widths, font_size=8.5, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for ri, row_values in enumerate(rows):
        new_row = table.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        if ri % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for i, value in enumerate(row_values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, bold=(first_col_bold and i == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def rec(family, scope, model, split):
    matches = [r for r in DATA if r["family"] == family and r["scope"] == scope and r["model"] == model and r["split"] == split]
    return matches[0] if matches else None


def pct(value):
    return "—" if value is None else f"{100 * value:.2f}"


def mean_sd(row, metric, sd_key):
    if not row or row.get(metric) is None:
        return "—"
    sd = row.get(sd_key)
    return pct(row[metric]) if sd is None else f"{pct(row[metric])} ± {100*sd:.2f}"


def add_split_perf_table(doc, family, scope, models):
    rows = []
    for model in models:
        row = [MODEL_NAMES.get(model, model)]
        for split in ("normal", "fault", "all"):
            rr = rec(family, scope, model, split)
            row.append("—" if rr is None else f"{pct(rr['node_accuracy'])} / {pct(rr['tier3_accuracy'])}")
        rows.append(row)
    add_table(doc, ["模型/配置", "Normal Node/Tier3", "Fault Node/Tier3", "All Node/Tier3"], rows,
              [3000, 2120, 2120, 2120], font_size=8.2, first_col_bold=True)
    add_para(doc, "表中单位为百分比，单元格为 Node Accuracy / Tier3 Accuracy。完整 Macro-F1 与 Balanced Accuracy 见随报告附带的 performance_summary.csv。", italic=True, color=MID_GRAY)


def add_all_metric_table(doc, family, scope, models):
    rows = []
    for model in models:
        rr = rec(family, scope, model, "all")
        rows.append([
            MODEL_NAMES.get(model, model),
            mean_sd(rr, "node_accuracy", "node_accuracy_sd"),
            mean_sd(rr, "node_macro_f1", "node_macro_f1_sd"),
            mean_sd(rr, "tier3_accuracy", "tier3_accuracy_sd"),
            mean_sd(rr, "tier3_macro_f1", "tier3_macro_f1_sd"),
        ])
    add_table(doc, ["模型/配置", "Node Acc", "Node Macro-F1", "Tier3 Acc", "Tier3 Macro-F1"], rows,
              [2850, 1600, 1800, 1500, 1610], font_size=8.0, first_col_bold=True)


doc = Document()
configure_document(doc)
add_cover(doc)

add_heading(doc, "1. 执行摘要", 1)
add_para(doc, "本报告回顾 Objective 3 中除实时动作边界分割和在线识别之外的全部主要实验。主线从 task graph 与 manifest 的可靠化开始，经 J-as-test 可行性先导、严格四折三 seed 基线、Direct Head Fusion、Dynamic Epoch Graph-Valid Shuffle、Atomic-tail 局部重排与真实顺序复评，最后进入 A0–A4 DualPos 消融。")
add_numbered(doc, [
    "最稳定、最清楚的核心收益来自历史位置消歧与 Direct Fusion：all-runs 的 M2 Direct 在 test_all 达到 90.57% Node Accuracy / 90.64% Tier3 Accuracy，相对 M0 为 +20.76 / +7.32 个百分点。",
    "固定 graph-valid shuffle 本身不是 Direct 模型的主要增益源。all-runs 中 M3 Direct 比 M2 Direct 低 0.52 pp Node Accuracy；normal-only 两者基本持平。",
    "每个 epoch 动态刷新合法拓扑序也没有提升：all-runs Dynamic Direct 为 89.79%，低于固定 M3 Direct 的 90.05%，更低于实际顺序 M2 Direct 的 90.57%。",
    "Atomic-tail 的意义在于只在活动中的未完成 atomic prefix 周围施加局部合法重排。统一用真实历史顺序测试后，all-runs 的 refresh-once 达到 91.03% / 91.18%，比 M2 Direct 高 +0.46 / +0.54 pp，但 12 个 fold×seed 配对为 6 胜 6 负，不能声称稳定取代 M2 Direct。",
    "A0–A3 消融解释了旧 Atomic-tail 不稳定性的来源：A1 broad shuffle 降到 89.35%，A2 active-tail-only 回升到 89.92%，A3 再保留 true recency 后达到 91.09%，说明 gating 与位置语义都重要。",
    "A3-DualPos 从头训练并未优于 A3（90.52%，-0.57 pp，2胜10负），且 Stage 3 下降 3.36 pp；显式位移不是自动有效的额外特征。",
    "A4-DualPos 从 A0 warm-start，以 2+8+3 epoch 的 shift 预热、paired 微调和 actual 校准，达到 91.02%，相对 A0 +0.45 pp，8胜1平3负，四个 participant 的三-seed 均值均为正；这是小而较稳定的工程提升，但不能全部归因于 DualPos。",
    "错误类型随模型演进发生迁移：M0 的 Node 错误中 45.8% 是同 Tier3 的重复流程位置混淆；M2 Direct 降到 0.8%，Atomic once actual-order 为 1.6%。后续瓶颈主要是物体、动作方向、设备开关状态与短时语义边界，而不是流程位置。",
])

add_heading(doc, "2. 报告边界、术语与统计口径", 1)
add_para(doc, "明确排除：realtime_action_boundary_experiment_2026-08-07 与 task_graph_realtime_demo_A_run7_2026-07-27 中的连续视频边界检测、因果滑窗和在线显示实验。本文只讨论预切分 clip 的离线识别、历史融合、图约束与训练增强。")
add_table(doc, ["项目", "统一定义"], [
    ["任务", "35 个 task-graph node 分类；node 概率进一步聚合为 31 类 Tier3"],
    ["输入", "RGB，camera 001484412812；每个 clip 16 帧、224×224；R3D-18 提取 512-D 特征"],
    ["严格外层协议", "A/D/J/M participant-level LOSO；每折仅用另外三人训练"],
    ["随机种子", "正式主线为 1、2、42；早期 J 先导为 1、2、3"],
    ["Train scope", "normal_only：只用正常 run 训练；all_runs：正常与故障 run 均进入训练"],
    ["Test split", "test_normal、test_fault、test_all；每个最终 checkpoint 都评估三者"],
    ["Checkpoint", "最后一个 epoch 的 last.pth；无 validation、无 early stopping、无 test-based 选模"],
    ["正式聚合", "先在每位 participant 内平均 3 seeds，再对 A/D/J/M 等权；± 为 4 位 participant 均值的样本 SD"],
    ["A0–A4 报告口径", "12 个 fold×seed 等权均值与样本 SD；因每人 seeds 数相同，均值与 participant-first 相同，SD 口径不同"],
], [1900, 7460], font_size=9.0, first_col_bold=True)
add_para(doc, "测试规模：A 431、D 462、J 555、M 447 个 clips，合计 1,895；其中 normal 1,441、fault 454。四人共有 103 个测试 runs（76 normal、27 fault）。D 的 fault 仅 62 clips 且类别覆盖不完整，因此 fault Macro-F1 的跨人比较需结合 support 解读。")

add_heading(doc, "3. 实验演进时间线与完成状态", 1)
timeline = [
    ["0", "2026-07-15", "Task graph/manifest 可靠化", "图结构审查；1,895 条主 manifest 与 15,614 条子 manifest 加 node/stage；无模型性能"],
    ["1", "2026-07-20", "J-as-test 先导", "M0–M6；existing Tier3 backbone；seeds 1/2/3；用于可行性，不作为严格四折结论"],
    ["2", "2026-07-22–27", "严格四折主网格", "A/D/J/M × 3 seeds × 2 scopes；M0–M6 + 3 E2E；完整"],
    ["3", "2026-07-29", "Direct Head Fusion", "M1/M2/M3 Direct；216 个 model-split 结果；完整"],
    ["4", "2026-07-30", "Dynamic epoch shuffle", "Frozen-M0 / Joint-Head / Direct；216 个 model-split 结果；完整"],
    ["5", "2026-08-03–04", "Atomic-tail refresh 网格", "Direct 分支；2 scopes × 3 refresh；72 个训练任务；完整"],
    ["6", "2026-08-05", "Atomic 真实顺序复评", "72 checkpoints × 3 splits = 216；不重训，只改测试历史顺序；完整"],
    ["7", "2026-08-19–20", "A0–A3 机制消融", "A0/A1/A2/A3；all_runs；4 folds × 3 seeds；完整"],
    ["8", "2026-08-20", "DualPos 与保守微调", "A3-DualPos、A4-DualPos；各 12/12；完整"],
    ["规划", "未运行", "A3-full、旧 A4、A5–A8", "deferred；只有设计与配置，不能报告性能"],
]
add_table(doc, ["阶段", "日期", "实验族", "完成内容"], timeline, [800, 1400, 2250, 4910], font_size=8.5, first_col_bold=True)

add_heading(doc, "4. 阶段 0：Task graph 与数据标注基础", 1)
add_para(doc, "这一阶段不是模型实验，但决定了后续 graph-valid 与 atomic-tail 是否可信。首先对 task_graph/task_graph.json 生成交互图、SVG/PNG 与结构/语义验证报告，源图不被自动修改。随后将 node_id、node_idx、stage_id 写入数据 manifest 的副本，并对重复 Tier3 标签使用 participant、run、annotation_row_index、Stage 2 序列及相邻转移做序列对齐。")
add_bullets(doc, [
    "主 manifest 输入/输出均为 1,895 条；unmatched labels=0，unresolved duplicate labels=0。",
    "20 个子 manifest 合计 15,614 行全部找到匹配；原文件有逐字节备份，更新后 hash 与验证副本一致。",
    "故障 run 中观测到 6 次 node 22→16 的重复动作转移，但标准 task graph 不加入该边；正常工艺图仍代表标准序列，故障偏离保留为异常。",
    "这一步建立 35-node、31-Tier3、3-stage 的统一标签空间，并为 graph-valid shuffle、relation matrix、active atomic prefix 判定提供可靠基础。",
])

add_heading(doc, "5. 阶段 1：J-as-test 可行性先导", 1)
add_para(doc, "目的：先在 J 作为 held-out participant 的条件下验证历史 attention、位置编码、graph-valid 重排与 relation bias 能否运行并产生方向性信号。该阶段复用了既有 Tier3 last.pth，且旧 backbone 训练曾将 J manifest 作为 validation；因此只能作为方法/工程先导，不能与后续 scratch-per-fold 严格结果混合。")
add_para(doc, "固定配置：normal-only 训练；seeds 1/2/3；M0–M6；同 run 因果历史；最后 epoch；每个模型评估 normal/fault/all。")
pilot_models = ["m0", "m1", "m2", "m3", "m4", "m5", "m6"]
add_split_perf_table(doc, "j_pilot_existing_backbone", "normal_only", pilot_models)
add_para(doc, "先导信号非常强：M3 的三-seed test_all 均值约 90% Node Accuracy，而 M0 约 74%；M2/M3 明显优于无位置 M1，说明历史的关键不是简单加入过去特征，而是用位置/流程结构消除重复动作 node 的歧义。后续阶段的目标是用严格四折、独立 scratch backbone 检验这一信号是否普遍。")

add_heading(doc, "6. 阶段 2：严格四折 M0–M6 与 E2E 对照", 1)
add_heading(doc, "6.1 实验含义与配置", 2)
add_para(doc, "每个 participant、seed、scope 都从 scratch 训练独立 R3D-18（100 epochs，batch 16，AdamW，LR 1e-4，epoch 50/75 衰减，最后 epoch），再冻结并缓存 512-D feature。Feature-level 模型训练 50 epochs，batch 64，AdamW，LR 1e-3，weight decay 1e-4。")
add_table(doc, ["模型", "回答的问题"], [
    ["M0", "只看当前 clip 的冻结视觉特征能做到什么？"],
    ["M1", "加入同 run 历史但不编码位置，单纯 history content 是否有用？"],
    ["M2", "真实时间顺序 + recency position 是否解决重复流程 node？"],
    ["M3", "将历史改为 task-graph 合法的固定拓扑重排，能否比实际顺序更好？"],
    ["M4", "35 个 candidate query 的结构本身能否提升？"],
    ["M5", "使用历史真实 node 的 oracle relation bias 的上限是多少？不可部署。"],
    ["M6", "用冻结 M0 历史概率形成 soft relation bias，是否得到可部署收益？"],
    ["3 个 E2E", "从 RGB 直接训练 Tier3/Node，或从 Tier3 迁移到 Node，作为视频端对照。"],
], [1800, 7560], font_size=8.8, first_col_bold=True)

main_models = ["m0", "m1", "m2", "m3", "m4", "m5", "m6", "e2e_node_scratch", "e2e_node_from_tier3", "e2e_tier3_scratch"]
add_heading(doc, "6.2 Normal-only 结果", 2)
add_split_perf_table(doc, "original_and_e2e", "normal_only", main_models)
add_all_metric_table(doc, "original_and_e2e", "normal_only", main_models)
add_para(doc, "normal-only 下，M0 test_all 为 66.76% Node Accuracy；M1–M6 提升到约 78.6–79.8%。M6 的 79.76% 最高，但 M2/M3/M4/M5/M6 差距小。位置编码使 M2 显著优于 M1，而固定 graph-valid M3 与实际顺序 M2 的 test_all Node Accuracy 同为 79.30%，没有额外增益。")

add_heading(doc, "6.3 All-runs 结果", 2)
add_split_perf_table(doc, "original_and_e2e", "all_runs", main_models)
add_all_metric_table(doc, "original_and_e2e", "all_runs", main_models)
add_para(doc, "all-runs 下，最强 delta 模型是 M3：84.74% Node Accuracy / 85.63% Tier3 Accuracy，相对 M0 的 69.81% / 83.32% 为 +14.94 / +2.31 pp。M3 比 M2 高 +0.59 pp Node、+0.64 pp Tier3，说明固定合法重排在 delta 架构中有小幅均值收益；但 relation M5/M6 只相对 M4 小幅改善，仍未稳定超过 M3。")
add_para(doc, "训练范围影响：all-runs 对历史模型通常有益，尤其 M3 相对 normal-only 的 test_all Node Accuracy 提高约 5.44 pp；但这同时改变 backbone 与下游训练数据，不能仅归因于 fault 样本本身。")

add_heading(doc, "7. 阶段 3：Direct Head Fusion", 1)
add_heading(doc, "7.1 为什么改架构", 2)
add_para(doc, "原 M1–M6 冻结 M0 node head，只学习一个 history delta 去修正 M0 logits。Direct 方案仍冻结 RGB backbone，却让 history fusion 与新的 35-node 分类头共同训练：current 512-D 与 attention context 256-D 拼接，经 768→512 fusion，再直接分类。Fusion 初始化为 current identity、history zero，避免训练初期随机 history 破坏视觉特征。")
add_para(doc, "M1/M2/M3 Direct 分别对应：无位置、真实顺序+位置、固定 graph-valid+位置。公共训练为 50 epochs、batch 64、LR 1e-3；A/D/J/M × 3 seeds × 2 scopes × 3 models × 3 splits = 216 个结果，网格完整。")
direct_models = ["m1_direct", "m2_direct", "m3_direct"]
add_heading(doc, "7.2 Normal-only", 2)
add_split_perf_table(doc, "direct_head_fusion", "normal_only", direct_models)
add_all_metric_table(doc, "direct_head_fusion", "normal_only", direct_models)
add_heading(doc, "7.3 All-runs", 2)
add_split_perf_table(doc, "direct_head_fusion", "all_runs", direct_models)
add_all_metric_table(doc, "direct_head_fusion", "all_runs", direct_models)
add_bullets(doc, [
    "all-runs M2 Direct 为 90.57% Node / 90.64% Tier3；相对 M0 为 +20.76 / +7.32 pp，12/12 fold×seed 配对均提高。",
    "M2 Direct 相对 M2 delta 为 +6.41 pp Node / +5.65 pp Tier3；联合训练分类头是大幅提升的核心。",
    "位置编码是关键：M2 Direct 相对 M1 Direct 为 +10.57 pp Node / +5.67 pp Tier3，12/12 配对均提高。",
    "固定 graph-valid 不是额外收益：all-runs M3 Direct 比 M2 Direct 低 0.52 pp Node / 0.37 pp Tier3；normal-only 只高 0.08 / 0.12 pp，差异可忽略。",
    "M2 Direct 相对 M0 在 103/103 个测试 run 上提高；重复动作 node 的四组双向误判由 M0 的 769 次降到 4 次。",
])

add_heading(doc, "8. 阶段 4：Dynamic Epoch Graph-Valid Shuffle", 1)
add_heading(doc, "8.1 实验含义与配置", 2)
add_para(doc, "原 M3/M3 Direct 在 Dataset 初始化时为每个样本生成一次合法拓扑序，50 epochs 内固定。Dynamic 实验每个 epoch、每个样本用 SHA256(base_seed:epoch:sample_name) 重新生成可复现的合法顺序；主测试仍使用与原 M3 相同的固定 seeded graph-valid 顺序，保证配对公平。")
add_table(doc, ["Dynamic 模型", "结构目的"], [
    ["Frozen-M0 Delta", "只改变 fixed→epoch-wise shuffle，最直接检验动态增强是否优于 M3"],
    ["Joint-Head Delta", "不加载 M0，新的 node head 与 delta/history 一起训练，检验冻结 head 的限制"],
    ["Direct Fusion", "动态重排条件下的 Direct 版本，检验动态增强是否能超过固定 M3 Direct"],
], [2500, 6860], font_size=9, first_col_bold=True)
dynamic_models = ["m3_dynamic_frozen_m0_delta", "m3_dynamic_joint_head_delta", "m3_dynamic_direct_fusion"]
add_heading(doc, "8.2 Normal-only", 2)
add_split_perf_table(doc, "dynamic_epoch_shuffle", "normal_only", dynamic_models)
add_all_metric_table(doc, "dynamic_epoch_shuffle", "normal_only", dynamic_models)
add_heading(doc, "8.3 All-runs", 2)
add_split_perf_table(doc, "dynamic_epoch_shuffle", "all_runs", dynamic_models)
add_all_metric_table(doc, "dynamic_epoch_shuffle", "all_runs", dynamic_models)
add_bullets(doc, [
    "all-runs Dynamic Frozen-M0 Delta 比固定 M3 低 0.51 pp Node / 0.50 pp Tier3。",
    "Dynamic Joint-Head 比 Dynamic Frozen 提高 1.14 / 0.93 pp，继续证明冻结 M0 head 是瓶颈。",
    "Dynamic Direct 达到 89.79 / 90.02%，比固定 M3 Direct 低 0.26 / 0.26 pp，也低于实际顺序 M2 Direct。",
    "逐样本比较中新增纠正少于新增退化；因此不能把更多合法排列暴露解释为有效数据增强。",
    "失效集中在 D 的低总体表现、A 的三-seed 一致错误，以及 node 1、34、4、8、24；剩余错误已主要是跨 Tier3 视觉语义混淆。",
])

add_heading(doc, "9. 阶段 5–6：Atomic-tail 局部 Graph-Valid 增强", 1)
add_heading(doc, "9.1 Atomic-tail 是什么", 2)
add_para(doc, "Atomic-tail 不再广泛重排全部历史。它查看最后一个真实历史 node，若其所属 atomic sequence 在已观察历史中恰好形成从首节点开始、尚未完成的 proper prefix，则把这一活动 tail 固定在历史尾部，只对其余历史做 graph-valid 重排。重复节点、非 prefix、完整 sequence 或图约束冲突均安全回退；重排函数不接收当前 target 或未来 clip。")
add_para(doc, "刷新频率比较 1、10、once；训练结构只完成 Direct Fusion 分支。原配置文件还定义 Frozen-M0 与 Joint-Head，但实际正式结果没有这两类模型，报告不把它们当成已完成实验。")
add_heading(doc, "9.2 原固定 Atomic 测试与实际顺序复评", 2)
add_para(doc, "最初训练和测试均使用 fixed seeded atomic-tail 顺序。为排除测试协议混杂，2026-08-05 对全部 72 个 checkpoint 不重训地改用 actual_chronological history 复评，生成 216 个 split 结果；旧结果未覆盖。两种测试顺序的 test_all Node 变化全部不超过 0.19 pp，说明结论不是由测试重排人为制造。以下以统一实际顺序结果为准。")
atomic_models = ["refresh_every_1", "refresh_every_10", "refresh_once"]
add_heading(doc, "9.3 Normal-only（actual-order evaluation）", 2)
add_split_perf_table(doc, "atomic_tail_refresh_grid_actual_eval", "normal_only", atomic_models)
add_heading(doc, "9.4 All-runs（actual-order evaluation）", 2)
add_split_perf_table(doc, "atomic_tail_refresh_grid_actual_eval", "all_runs", atomic_models)
add_bullets(doc, [
    "all-runs 最佳为 refresh-once：normal 91.69/91.80、fault 90.19/90.37、all 91.03/91.18；均比 M2 Direct 高约 0.4–0.5 pp。",
    "normal-only 最佳为 every-10：normal 91.20/91.39、fault 84.28/85.75、all 89.42/89.85。最优刷新频率依赖训练范围。",
    "每 epoch 刷新在两个 scope 都最差或接近最差，说明扰动过频会破坏稳定表示。",
    "all-runs once 相对 M2 Direct 的 12 个配对为 6胜6负，Node 描述性 95% 区间 [-1.45,+2.37] pp；A/D/J 提升，M 下降 0.89 pp。",
    "normal-only every-10 为 6胜1平5负，A 下降、J/M 提升；同样不能声称跨人稳定。",
    "72 份 shuffle audit 的 atomic_tail_violations 总数为 0，说明所有生成历史保持约束合法。",
])

add_heading(doc, "10. 阶段 7：A0–A3 Atomic-tail 机制消融", 1)
add_heading(doc, "10.1 为什么重新做 A0–A8 包", 2)
add_para(doc, "旧 refresh 网格显示 Atomic-tail 均值有潜力但 seed/participant 稳定性不足。新包将测试统一为实际顺序，并把 broad shuffle、active-tail gating、位置语义、DualPos、warm-start/paired training 等因素逐步拆开。固定条件为 all_runs、A/D/J/M、seeds 1/2/42、Direct Fusion、512-D 冻结特征、50-epoch scratch（除 A0/A4-DualPos）。")
add_table(doc, ["实验", "唯一/主要变化", "状态"], [
    ["A0", "实际顺序 M2 Direct；复用匹配的共享 checkpoint", "12/12 完整"],
    ["A1", "旧 broad atomic-tail once；无 active tail 也全历史 graph-valid shuffle；presented position", "12/12 完整"],
    ["A2", "只允许 active incomplete tail 样本增强；presented position", "12/12 完整"],
    ["A3", "A2 + 每个历史元素保留 actual recency position", "12/12 完整"],
    ["A3-full-shuffle", "A1 的 broad scope + A3 的 true recency；在 single-query attention 下扰动近似不可见", "deferred"],
], [1200, 6260, 1900], font_size=8.8, first_col_bold=True)
a_models = ["A0", "A1", "A2", "A3", "A3-DualPos", "A4-DualPos"]
add_split_perf_table(doc, "atomic_tail_A0_A4_dualpos", "all_runs", a_models)
add_all_metric_table(doc, "atomic_tail_A0_A4_dualpos", "all_runs", a_models)
add_bullets(doc, [
    "A1 比 A0 低 1.22 pp Node Accuracy：broad shuffle + presented position 会引入不必要的训练/测试不一致。",
    "A2 比 A1 回升 0.57 pp：active-tail-only gating 减少无关扰动，但仍低于 A0 0.65 pp。",
    "A3 比 A2 再提高 1.17 pp，并比 A0 高 0.52 pp：让事件携带真实 recency 是关键修正。",
    "A3 的 normal 为 91.85%，fault 为 89.78%，all 为 91.09%；提升主要不来自 fault 专项，而是更合理的历史位置语义。",
    "训练 audit 显示 active tail 约覆盖 69.39% 样本，但真正改变顺序的样本仅 17.79%；增强是稀疏、局部的。",
])

add_heading(doc, "11. 阶段 8：DualPos 与 A4-DualPos", 1)
add_heading(doc, "11.1 DualPos 定义", 2)
add_para(doc, "A3 在 single-query attention 中把每个历史 token 的特征与真实 recency 绑定后，单纯改变 token 呈现顺序近似是集合置换，模型可能看不见 shuffle。DualPos 因此同时编码真实 recency r 与增强后呈现位置 p 的位移 Δ=p−r：token = Wf·x + E_true(r) + E_shift(Δ)。测试为实际顺序，Δ=0 的 padding row 固定为零，因此新增分支不直接改变 A0 推理路径。")
add_table(doc, ["实验", "训练日程", "要检验的问题"], [
    ["A3-DualPos", "从头 50 epochs；active-tail true-recency + shift；固定 once shuffle", "位移显式可见本身是否改善 A3？"],
    ["A4-DualPos", "A0 warm-start；2 epoch shift-only + 8 paired joint + 3 actual calibration", "把 DualPos 作为强基线附近的局部正则是否更稳定？"],
], [1800, 3550, 4010], font_size=8.8, first_col_bold=True)
add_bullets(doc, [
    "A3-DualPos 为 90.52%，比 A3 低 0.57 pp；12 个配对 2胜10负。Stage 1 +0.91 pp，但 Stage 3 -3.36 pp，说明位移信号在长历史/后段流程可能过强或难优化。",
    "A4-DualPos 为 91.02%，相对 A0 +0.45 pp；8胜1平3负，Node 配对 95% CI [+0.01,+0.89] pp。四个 participant 的三-seed均值均为正。",
    "A4-DualPos 与 A3 基本同水平（-0.07 pp Node、+0.06 pp Node Macro-F1），但相对 A3 的 Stage 3 仍低 1.55 pp。",
    "A4-DualPos 与 A0 的预测一致率 97.54%；5,685 个 clip-seed 预测中净增加 26 个正确，run-cluster bootstrap 95% CI [+0.10,+0.79] pp。它是保守局部修正，而非大范围改变决策边界。",
    "DualPos audit：15.90% history tokens 有非零 shift，全部 token 的 mean absolute shift=0.556；按 shifted tokens 条件化约 3.49 个位置。代码确实施加了位移，A3-DualPos 的下降不是未生效。",
    "A4-DualPos 同时改变 warm-start、训练日程、paired view、refresh 与 calibration，缺少相同日程的 A4-NoShift，故 +0.45 pp 不能全部归因于位移 embedding。",
])

add_heading(doc, "12. 跨阶段综合：我们真正学到了什么", 1)
add_table(doc, ["问题", "证据", "当前判断"], [
    ["历史是否有效？", "M2 Direct vs M0：+20.76 pp Node；103/103 runs 提升", "强支持"],
    ["位置编码是否关键？", "M2 Direct vs M1 Direct：+10.57 pp；12/12 配对提升", "强支持"],
    ["固定 graph-valid 是否优于实际顺序？", "M3 Direct 比 M2 Direct -0.52 pp；delta 中仅 +0.59 pp", "不稳定/架构依赖"],
    ["动态 shuffle 是否有效？", "Dynamic Direct 比固定 M3 Direct -0.26 pp", "不支持"],
    ["Atomic-tail 是否有潜力？", "actual-eval once 91.03%，高 M2 Direct +0.46 pp，但 6胜6负", "小幅均值优势，不稳定"],
    ["active-tail gating 是否重要？", "A2 比 A1 +0.57 pp", "支持"],
    ["true recency 是否重要？", "A3 比 A2 +1.17 pp", "支持"],
    ["DualPos 单独是否有效？", "A3-DualPos 比 A3 -0.57 pp，2胜10负", "不支持"],
    ["保守 A4-DualPos 是否有用？", "A4-DualPos 比 A0 +0.45 pp，8胜1平3负", "小幅正向，机制未拆清"],
], [2350, 4300, 2710], font_size=8.4, first_col_bold=True)

add_heading(doc, "13. 失效模式与性能上限", 1)
add_para(doc, "模型的错误不再是同一种错误。M0 的 1,698 个 Node 错误中 45.8% 是同一 Tier3 的重复流程 node 混淆；四组重复 node 双向错误共 769 次。M3 降至 5.7%/46 次，M2 Direct 降至 0.8%/4 次，Atomic once actual-order 为 1.6%/7 次。流程位置问题基本被历史解决，剩余约 99% 错误转为跨 Tier3 视觉语义错误。")
add_bullets(doc, [
    "相同空间区域、不同物体：node 24 put sample on table 常错为 node 12 take plier、25 put plier 或 34 take lock。",
    "相同物体、相反动作方向：put sample 与 grip sample（16/19、17/20）共享物体和区域，只靠短 clip 难辨运动方向。",
    "设备状态相反：6↔30、7↔29、8↔28 及 lock/unlock 的外观变化短暂。",
    "短时语义边界：node 18 reverse 与 23 inspect，或 place sample 与 press pedal，clip 可能包含相邻动作过渡。",
    "Participant 特异：A 主要难在 node 24，D 在 node 18，J 在 node 34，M 在 setup/shutdown；单一增强策略难以同时改善所有人。",
    "高置信错误仍多：M2 Direct、Dynamic Direct、Atomic once 的高置信错误比例约 48.1%、46.4%、46.8%，softmax 最大值不能可靠识别失败。",
])

add_heading(doc, "14. 未完成实验与下一步优先级", 1)
add_table(doc, ["配置", "当前状态", "继续前需要回答"], [
    ["A3-full-shuffle", "deferred", "在 true-recency single-query attention 下近似 token 集合置换，预期信息量低"],
    ["旧 A4（无 shift）", "deferred", "paired actual/aug 在当前 attention 下近似同一集合；不建议按旧定义直接运行"],
    ["A5 consistency", "deferred", "需先迁移到 DualPos 方案并确认 consistency weight/threshold"],
    ["A6 plausibility sampling", "deferred", "需审核 candidate changed fraction、Kendall distance 与训练折转移先验"],
    ["A7 tail-order auxiliary", "deferred", "需报告 eligible fraction，防止因有效样本太少误判"],
    ["A8 Tier3 auxiliary", "deferred", "需同时报告 Node 与 Tier3，避免只选择更好看的层级指标"],
], [1900, 1500, 5960], font_size=8.6, first_col_bold=True)
add_numbered(doc, [
    "最高优先级：A4-DualPos-NoShift。完全复制 A4-DualPos 的 A0 checkpoint、2+8+3 日程、paired weights、refresh 和优化器，只强制 history_shift_ids=0；这是把 +0.45 pp 归因给 DualPos 的必要控制。",
    "零重训优先：直接评估 A4-DualPos 已保存的 after_shift_warmup、after_mixed_finetune、after_actual_calibration 三个 checkpoints，判断增益来自 paired 阶段还是最后 calibration。",
    "方法层面：为 shift embedding 加小初值 gate 或随 history length 衰减，重点缓解 Stage 3 退化。",
    "视觉层面：优先加入 hand-object crop/object token、动作前后状态差、较长 clip 或 motion cue；当前瓶颈已不主要是 graph 顺序。",
    "报告层面：同时给出 M2 Direct 与最高均值 Atomic/A4-DualPos，保留 fold×seed 配对、participant 差异、support 与不确定性，不把 clip 当作独立受试者。",
])

add_heading(doc, "15. 推荐用于论文或阶段汇报的表述", 1)
add_para(doc, "可使用：历史位置编码与 task graph 上下文显著改善了跨人 35-node 流程识别，其主要作用是消除视觉相同动作在不同流程位置的歧义；在冻结视觉 backbone 的条件下，history fusion 与新的 node head 联合训练明显优于冻结 M0 后仅学习 logit delta。")
add_para(doc, "可使用但需限定：Atomic-tail/true-recency 训练增强在统一实际顺序测试下获得约 0.5 pp 的最高均值，并在 A0–A3 消融中显示 active-tail gating 与真实 recency 有正向作用；然而收益随 participant 与 seed 变化，当前尚未稳定显著超过 M2 Direct。")
add_para(doc, "暂不使用：‘graph-valid shuffle 普遍提升准确率’、‘每 epoch 动态重排带来更强泛化’、‘DualPos 本身导致 A4-DualPos 的全部收益’。现有严格消融不支持这些表述。")

add_heading(doc, "16. 结果来源与可复核文件", 1)
sources = [
    ["主实验配置", "codex_and_files/graph_history_rgb_cross_person_ADM_2026-07-22/COMPLETE_EXPERIMENT_CONFIGURATION.md"],
    ["主线完整分析", ".../EXPERIMENT_RESULTS_ANALYSIS_2026-08-04.md（更新至 2026-08-05）"],
    ["M0–M6/E2E", ".../outputs/cross_person_summary_*_ADJM_3seeds/all_model_cross_person_aggregate.csv"],
    ["Direct", ".../outputs/direct_head_fusion_summary_ADJM_3seeds/direct_head_aggregate.csv"],
    ["Dynamic", ".../outputs/dynamic_epoch_shuffle_summary_ADJM_3seeds/dynamic_epoch_shuffle_aggregate.csv"],
    ["Atomic actual-eval", ".../outputs/at_actual/atomic_actual_order_aggregate.csv"],
    ["A0–A8 配置", "codex_and_files/atomic_tail_A0_A8_windows_2026-08-19/EXPERIMENT_CONFIGURATION.md"],
    ["A0–A4 分析", ".../A0_A4_DUALPOS_RESULTS_ANALYSIS_REPORT.md"],
    ["本报告机器表", "non_realtime_experiment_summary_2026-08-24/performance_summary.csv 与 .json"],
]
add_table(doc, ["内容", "来源"], sources, [2200, 7160], font_size=8.2, first_col_bold=True)
add_para(doc, "证据优先级：原始 metrics/predictions 与 completed/audit > 汇总 CSV > 已有分析报告 > 配置草案。报告将 deferred 配置与实测结果严格分开。")

add_heading(doc, "附录 A：快速选型", 1)
add_table(doc, ["使用场景", "推荐配置", "理由"], [
    ["最简单、最稳妥主模型", "all-runs M2 Direct", "实际顺序；90.57% Node；强配对收益；机制最清楚"],
    ["报告最高历史均值", "Atomic all-runs refresh-once actual-eval", "91.03% Node；但需同时报告 6胜6负"],
    ["新包最高 Node 均值", "A3", "91.09%；true recency 有机制证据；seed 交互较强"],
    ["更保守的小幅增益", "A4-DualPos", "91.02%；相对 A0 8胜1平3负；需 no-shift 控制"],
    ["不推荐", "every-epoch Dynamic/Atomic、A3-DualPos scratch", "严格结果未带来总体提升或出现退化"],
], [2600, 2900, 3860], font_size=8.5, first_col_bold=True)

doc.core_properties.title = "Objective 3 非 real-time 实验全流程总结"
doc.core_properties.subject = "Graph-history, graph-valid shuffle, atomic-tail and DualPos experiments"
doc.core_properties.author = "Codex"
doc.core_properties.comments = "Generated from project experiment outputs on 2026-08-24"
doc.save(DOCX)
print(DOCX)
