from __future__ import annotations

import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA = json.loads((ROOT / "analysis_data.json").read_text(encoding="utf-8"))
OUT = ROOT / "Sequence_Disjoint_Atomic_Tail_实验分析报告_2026-08-25.docx"
ASSET_DIR = ROOT / "report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

EXPERIMENTS = [
    "M2-Direct-RealOrder", "A1-Legacy-Once", "A1-Legacy-Every10-Replace",
    "A3-DualPos-Once", "A3-DualPos-Every10",
]
SHORT = {
    "M2-Direct-RealOrder": "M2 RealOrder",
    "A1-Legacy-Once": "A1 Once",
    "A1-Legacy-Every10-Replace": "A1 Every10",
    "A3-DualPos-Once": "DualPos Once",
    "A3-DualPos-Every10": "DualPos Every10",
}
PARTICIPANTS = ["A", "D", "J", "M"]
SEEDS = [1, 2, 42]
SPLITS = ["normal", "fault", "all"]
SPLIT_CN = {"normal": "Normal", "fault": "Fault", "all": "All"}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GOLD = "7A5A00"
GREEN = "1F5D42"
RED = "9B1C1C"
WHITE = "FFFFFF"
MID_GRAY = "68717A"
GRID = "C8D0D8"
FONT_CJK = "Microsoft YaHei"
FONT_LATIN = "Calibri"


def read_font(size: int, bold: bool = False):
    path = Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc")
    return ImageFont.truetype(str(path), size=size)


def pct(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}"


def mean_sd(entry: dict) -> str:
    return f"{pct(entry['mean'])} ± {pct(entry['sd'])}"


def set_run_font(run, size=None, bold=None, italic=None, color=None, latin=FONT_LATIN, east_asia=FONT_CJK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def cell_text(cell, text, *, bold=False, size=8.5, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=8.5, first_col_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT_GRAY)
        cell_text(table.rows[0].cells[i], header, bold=True, size=font_size, color=DARK_BLUE,
                  align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_header(table.rows[0])
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        if ridx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "FAFBFC")
        for i, value in enumerate(row_data):
            cell_text(row.cells[i], value, size=font_size,
                      align=WD_ALIGN_PARAGRAPH.LEFT if (i == 0 and first_col_left) else WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=9, italic=True, color=MID_GRAY)


def add_paragraph(doc, text="", *, bold_prefix=None, italic=False, color=INK, align=None, after=6, before=0, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if keep:
        p.paragraph_format.keep_with_next = True
    if bold_prefix and text.startswith(bold_prefix):
        a = p.add_run(bold_prefix)
        set_run_font(a, bold=True, color=color)
        b = p.add_run(text[len(bold_prefix):])
        set_run_font(b, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_callout(doc, label, text, tone="blue"):
    colors = {"blue": (LIGHT_BLUE, DARK_BLUE), "gold": ("FFF8E8", GOLD), "green": ("ECF7F1", GREEN), "red": ("FCEEEE", RED)}
    fill, ink = colors[tone]
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=ink)
    r = p.add_run(text)
    set_run_font(r, color=INK)
    set_table_geometry(table, [9360])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("第 ")
    set_run_font(r, size=9, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    r = paragraph.add_run(" 页")
    set_run_font(r, size=9, color=MID_GRAY)


def setup_document(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    heading_tokens = {
        1: (16, BLUE, 16, 8), 2: (13, BLUE, 12, 6), 3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("OBJECTIVE 3  |  SEQUENCE-DISJOINT EXPERIMENT")
    set_run_font(r, size=8.5, bold=True, color=MID_GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def chart_overall_accuracy(path):
    w, h = 1700, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title, label, small = read_font(38, True), read_font(24), read_font(20)
    d.text((70, 35), "Tier3 Accuracy：五种历史模型的总体比较", font=title, fill="#0B2545")
    left, top, right, bottom = 120, 130, 1650, 760
    ymin, ymax = 0.65, 0.90
    for i in range(6):
        yval = ymin + i * (ymax - ymin) / 5
        y = bottom - (yval - ymin) / (ymax - ymin) * (bottom - top)
        d.line((left, y, right, y), fill="#E2E7EC", width=2)
        d.text((35, y - 13), f"{yval*100:.0f}%", font=small, fill="#68717A")
    colors = ["#4E79A7", "#59A14F", "#8CD17D", "#E15759", "#F28E2B"]
    group_w = (right - left) / 3
    bar_w = 58
    gap = 12
    for si, split in enumerate(SPLITS):
        center = left + group_w * (si + 0.5)
        start = center - (5 * bar_w + 4 * gap) / 2
        for ei, exp in enumerate(EXPERIMENTS):
            value = DATA["aggregates"][exp][split]["tier3_accuracy"]["mean"]
            x0 = start + ei * (bar_w + gap)
            y0 = bottom - (value - ymin) / (ymax - ymin) * (bottom - top)
            d.rounded_rectangle((x0, y0, x0 + bar_w, bottom), radius=5, fill=colors[ei])
            d.text((x0 - 3, y0 - 28), f"{value*100:.1f}", font=small, fill="#263645")
        label_text = SPLIT_CN[split]
        box = d.textbbox((0, 0), label_text, font=label)
        d.text((center - (box[2] - box[0]) / 2, bottom + 18), label_text, font=label, fill="#0B2545")
    lx, ly = 210, 820
    for ei, exp in enumerate(EXPERIMENTS):
        d.rounded_rectangle((lx, ly, lx + 24, ly + 24), radius=4, fill=colors[ei])
        d.text((lx + 34, ly - 3), SHORT[exp], font=small, fill="#263645")
        lx += 290
    img.save(path)


def chart_delta_heatmap(path):
    w, h = 1500, 860
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title, label, small = read_font(38, True), read_font(25, True), read_font(23)
    d.text((60, 35), "相对 M2 的 Tier3 Accuracy 变化（百分点）", font=title, fill="#0B2545")
    x0, y0, cw, ch = 500, 160, 280, 130
    for j, split in enumerate(SPLITS):
        d.text((x0 + j*cw + 85, 115), SPLIT_CN[split], font=label, fill="#1F4D78")
    for i, exp in enumerate(EXPERIMENTS[1:]):
        d.text((60, y0 + i*ch + 42), SHORT[exp], font=label, fill="#263645")
        for j, split in enumerate(SPLITS):
            val = DATA["deltas_vs_m2"][exp][split]["tier3_accuracy"]["mean"] * 100
            strength = min(abs(val) / 3.0, 1.0)
            if val >= 0:
                base, target = (244, 250, 246), (51, 132, 91)
            else:
                base, target = (253, 245, 245), (181, 62, 62)
            color = tuple(int(base[k] + strength * (target[k] - base[k])) for k in range(3))
            xa, ya = x0 + j*cw, y0 + i*ch
            d.rounded_rectangle((xa, ya, xa+cw-18, ya+ch-18), radius=12, fill=color)
            text = f"{val:+.2f}"
            box = d.textbbox((0, 0), text, font=label)
            ink = "white" if strength > 0.55 else "#263645"
            d.text((xa+(cw-18-(box[2]-box[0]))/2, ya+38), text, font=label, fill=ink)
    d.text((60, 735), "绿色表示优于 M2，红色表示低于 M2；数值为12个 fold×seed 配对差值的均值。", font=small, fill="#68717A")
    img.save(path)


def chart_fold_all(path):
    w, h = 1600, 900
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title, label, small = read_font(38, True), read_font(25), read_font(20)
    d.text((65, 35), "All Split：按 held-out participant 的 Tier3 Accuracy", font=title, fill="#0B2545")
    left, top, right, bottom = 130, 145, 1525, 745
    ymin, ymax = 0.68, 0.96
    for i in range(8):
        value = ymin + i*(ymax-ymin)/7
        y = bottom - (value-ymin)/(ymax-ymin)*(bottom-top)
        d.line((left, y, right, y), fill="#E2E7EC", width=2)
        d.text((38, y-12), f"{value*100:.0f}%", font=small, fill="#68717A")
    colors = ["#4E79A7", "#59A14F", "#8CD17D", "#E15759", "#F28E2B"]
    xs = [left + i*(right-left)/3 for i in range(4)]
    for i, p in enumerate(PARTICIPANTS):
        d.text((xs[i]-10, bottom+20), p, font=label, fill="#0B2545")
    for ei, exp in enumerate(EXPERIMENTS):
        points = []
        for i, p in enumerate(PARTICIPANTS):
            value = DATA["fold_means"][exp]["all"][p]["tier3_accuracy"]
            y = bottom - (value-ymin)/(ymax-ymin)*(bottom-top)
            points.append((xs[i], y))
        d.line(points, fill=colors[ei], width=5)
        for x, y in points:
            d.ellipse((x-9, y-9, x+9, y+9), fill=colors[ei], outline="white", width=2)
    lx, ly = 180, 820
    for ei, exp in enumerate(EXPERIMENTS):
        d.line((lx, ly+10, lx+30, ly+10), fill=colors[ei], width=5)
        d.text((lx+40, ly-4), SHORT[exp], font=small, fill="#263645")
        lx += 275
    img.save(path)


def chart_coverage(path):
    w, h = 1600, 850
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title, label, small = read_font(38, True), read_font(23), read_font(20)
    d.text((65, 35), "增强强度与测试历史前缀覆盖（四折均值）", font=title, fill="#0B2545")
    exps = EXPERIMENTS[1:]
    left, top, right, bottom = 120, 150, 1530, 700
    for i in range(7):
        val = i*0.1
        y = bottom - val/0.6*(bottom-top)
        d.line((left, y, right, y), fill="#E2E7EC", width=2)
        d.text((40, y-12), f"{val*100:.0f}%", font=small, fill="#68717A")
    group_w = (right-left)/4
    colors = ["#4E79A7", "#B0B8C1", "#59A14F"]
    for i, exp in enumerate(exps):
        center = left + group_w*(i+0.5)
        vals = [
            DATA["coverage"][exp]["mean_changed_fraction"],
            DATA["coverage"][exp]["mean_actual_test_prefix_coverage"],
            DATA["coverage"][exp]["mean_augmented_test_prefix_coverage"],
        ]
        for j, val in enumerate(vals):
            x = center - 84 + j*62
            y = bottom - val/0.6*(bottom-top)
            d.rounded_rectangle((x, y, x+50, bottom), radius=5, fill=colors[j])
            d.text((x-2, y-27), f"{val*100:.1f}", font=small, fill="#263645")
        name = SHORT[exp].replace(" ", "\n", 1)
        d.multiline_text((center-82, bottom+18), name, font=small, fill="#0B2545", align="center", spacing=2)
    legends = [("changed", colors[0]), ("actual prefix coverage", colors[1]), ("augmented prefix coverage", colors[2])]
    x = 280
    for text, color in legends:
        d.rectangle((x, 805, x+24, 829), fill=color)
        d.text((x+34, 801), text, font=small, fill="#263645")
        x += 370
    img.save(path)


def make_charts():
    paths = {
        "overall": ASSET_DIR / "figure_1_overall_tier3_accuracy.png",
        "delta": ASSET_DIR / "figure_2_delta_vs_m2.png",
        "fold": ASSET_DIR / "figure_3_fold_all_accuracy.png",
        "coverage": ASSET_DIR / "figure_4_augmentation_coverage.png",
    }
    chart_overall_accuracy(paths["overall"])
    chart_delta_heatmap(paths["delta"])
    chart_fold_all(paths["fold"])
    chart_coverage(paths["coverage"])
    return paths


def add_image(doc, path, width=6.35, alt="实验结果图表"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", alt)
    picture._inline.docPr.set("title", alt)


def build_report():
    charts = make_charts()
    doc = Document()
    setup_document(doc)

    # Editorial-cover opening; named override: Chinese title stack and restrained blue accent.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验分析报告")
    set_run_font(r, size=12, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Sequence-Disjoint 条件下的\nGraph-Valid Shuffle 与 Atomic-Tail 实验")
    set_run_font(r, size=26, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(55)
    r = p.add_run("M2-Direct-RealOrder、A1-Legacy 与 A3-DualPos 的系统比较")
    set_run_font(r, size=13.5, color=DARK_BLUE)
    add_table(doc, ["项目", "内容"], [
        ["实验包", "atomic_tail_sequence_disjoint_2026-08-24"],
        ["评估协议", "4-fold cross-person LOSO；每折3个随机种子（1、2、42）"],
        ["完成规模", "12个 backbone/feature 任务 + 60个 history-model 任务"],
        ["报告日期", "2026-08-25"],
    ], [1800, 7560], font_size=9.5)
    add_paragraph(doc, "本报告仅分析已完成的 sequence-disjoint 实验；所有主结果均来自 actual chronological order 测试。",
                  italic=True, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=18)
    doc.add_page_break()

    add_heading(doc, "执行摘要", 1)
    add_callout(doc, "核心结论", "在严格移除训练集中与测试对象完全同序的真实 run 后，A1-Legacy-Once 取得最好的综合结果：Tier3 All Accuracy 为85.46%，比 M2-Direct-RealOrder 高1.58个百分点。", "green")
    add_callout(doc, "边界条件", "优势不是跨参与者一致的：A、D 两折受益，J、M 两折轻微下降；All 的12个配对结果为7胜、1平、4负。因此证据支持“平均上有帮助”，但不足以宣称对所有未知顺序稳定占优。", "gold")
    add_callout(doc, "刷新策略", "A1 Every10 虽覆盖了更多测试历史前缀，但 All Accuracy 比 A1 Once 低0.89个百分点；更多顺序多样性没有自动转化为更高泛化性能。", "blue")
    add_callout(doc, "DualPos", "A3-DualPos-Once 与 Every10 的 All Accuracy 均约82.6%，低于 M2 约1.2个百分点。由于它同时改变了增强范围和位置语义，本轮不能把差异单独归因于 DualPos 编码。", "red")
    add_paragraph(doc, "按平均性能排序，Normal 与 All 的最佳模型都是 A1-Legacy-Once；Fault 的最佳模型是 A1-Legacy-Every10-Replace。Fault 样本量更小且跨参与者波动更大，相关提升应谨慎解释。")

    add_heading(doc, "1. 研究问题与实验逻辑", 1)
    add_heading(doc, "1.1 要验证的主张", 2)
    add_paragraph(doc, "本轮实验希望回答：当测试 run 的完整动作顺序从未作为真实训练 run 出现时，graph-valid shuffle augmentation 是否能通过生成新的合法历史顺序，使历史融合模型优于只使用真实训练顺序的 M2-Direct baseline。")
    add_paragraph(doc, "这里的关键不是禁止增强结果碰到测试顺序。增强器可以自然生成与测试相同的顺序，因为这正是顺序增强希望获得的泛化覆盖；唯一禁止的是增强器读取、定位或针对测试顺序进行采样。")
    add_heading(doc, "1.2 公平比较原则", 2)
    add_paragraph(doc, "所有五个历史模型共享同一 fold×seed 下重新训练的视觉 backbone 与新提取的512维特征；M2 始终使用真实历史顺序且不 shuffle；四种增强模型只在训练视图上改变历史顺序；测试统一使用真实 chronological order。")

    add_heading(doc, "2. 数据隔离协议", 1)
    add_heading(doc, "2.1 顺序定义与过滤规则", 2)
    add_paragraph(doc, "每个 run 按 annotation_row_index 排序，并以完整 node_idx 元组作为签名；连续重复节点不折叠。若一个训练 run 的完整签名与该折 test_all 中任一 run 完全一致，则移除整个训练 run。Tier3 顺序只用于审计，不参与过滤。")
    rows = []
    for p in PARTICIPANTS:
        z = DATA["isolation"][p]
        rows.append([p, z["source_train_runs"], z["filtered_train_runs"], z["removed_runs"],
                     f"{z['removed_runs']/z['source_train_runs']*100:.1f}%", z["filtered_train_samples"],
                     f"{z['normal_test_samples']}/{z['fault_test_samples']}/{z['all_test_samples']}", z["remaining_exact_overlap"]])
    add_table(doc, ["Test fold", "原train runs", "保留runs", "移除runs", "移除比例", "保留clips", "N/F/All clips", "剩余重叠"], rows,
              [1050, 1080, 1020, 1020, 980, 1040, 1900, 1270], font_size=8.2)
    add_caption(doc, "表1  Sequence-disjoint 过滤规模。四折均保持35个Node与31个Tier3类别完整覆盖。")
    add_paragraph(doc, "过滤强度很大：每折移除了40–52个训练 run，保留27–38个 run。A fold 仅保留623个训练 clips，是数据缩减最明显的一折，也正是增强收益最大的一折。这个对应关系值得注意，但不能据此直接证明因果。")

    add_heading(doc, "3. 实验配置与模型含义", 1)
    add_heading(doc, "3.1 上游视觉模型与特征", 2)
    add_paragraph(doc, "每个 fold×seed 均在过滤后的 all-runs train manifest 上从头训练 R3D-18 Tier3 classifier。实际训练日志确认全部为100 epochs；batch size 16，AdamW，初始学习率1e-4，weight decay 1e-4，AMP开启；学习率在第50、75轮后各衰减10倍。随后使用 last checkpoint 分别提取过滤 train 与完整 test_all 的512-D特征。旧 checkpoint 与旧 feature cache 均未复用。")
    add_heading(doc, "3.2 历史融合模型", 2)
    add_paragraph(doc, "历史模型共同配置为：feature_dim 512、d_model 256、4个attention heads、最大历史35、dropout 0.1；35个Node输出并聚合为31个Tier3类别。每个模型从头训练50 epochs，batch size 64，AdamW，学习率1e-3，weight decay 1e-4，gradient clip 1.0，关闭AMP，确定性训练。")
    rows = [
        ["M2 RealOrder", "真实顺序", "presented", "无", "无", "历史信息+真实顺序基线；严禁shuffle"],
        ["A1 Once", "广义graph-valid", "先shuffle后赋presented位置", "非active-tail也可增强", "仅一次", "整个训练保持固定增强视图"],
        ["A1 Every10", "广义graph-valid", "先shuffle后赋presented位置", "非active-tail也可增强", "每10轮替换", "轮次0/1/2/3/4；模型、位置嵌入、优化器连续"],
        ["DualPos Once", "graph-valid", "真实recency+位移embedding", "仅active atomic tail", "仅一次", "保留真实时间距离并编码呈现位移"],
        ["DualPos Every10", "graph-valid", "真实recency+位移embedding", "仅active atomic tail", "每10轮替换", "同一模型连续学习多轮增强视图"],
    ]
    add_table(doc, ["模型", "训练顺序", "位置语义", "增强范围", "刷新", "实验含义"], rows,
              [1450, 1450, 2050, 1600, 1100, 1710], font_size=8.0)
    add_caption(doc, "表2  五种历史模型的核心差异。")

    add_heading(doc, "4. 完整性与运行审计", 1)
    hist_ok = sum(not j["missing"] and j["epochs"] == 50 for j in DATA["integrity"]["history_jobs"])
    upstream_ok = sum(not j["missing"] and j["backbone_epochs"] == 100 for j in DATA["integrity"]["upstream_jobs"])
    rows = [
        ["Sequence-disjoint folds", "4/4", "exact full-run overlap=0；类别支持完整"],
        ["Backbone checkpoints", f"{upstream_ok}/12", "每个fold×seed从头训练100轮"],
        ["Feature caches", "12/12", "train_all.pt与test_all.pt及metadata齐全"],
        ["History jobs", f"{hist_ok}/60", "5模型×4 folds×3 seeds；均50轮"],
        ["测试指标", "180/180", "60 jobs×normal/fault/all；actual chronological order"],
    ]
    add_table(doc, ["审计项", "结果", "说明"], rows, [2300, 1500, 5560], font_size=9.2)
    add_callout(doc, "配置核对", "实验包配置与实际日志一致地显示 backbone=100 epochs。本报告采用实际训练日志作为最终事实来源。", "blue")

    add_heading(doc, "5. 上游 Backbone 性能", 1)
    rows = []
    for split in SPLITS:
        z = DATA["backbone_aggregates"][split]
        rows.append([SPLIT_CN[split], mean_sd(z["tier3_accuracy"]), mean_sd(z["tier3_macro_f1"]), mean_sd(z["tier3_balanced_accuracy"])])
    add_table(doc, ["Split", "Tier3 Accuracy", "Macro-F1", "Balanced Accuracy"], rows,
              [1800, 2520, 2520, 2520], font_size=9.2)
    add_caption(doc, "表3  R3D-18 last-checkpoint 性能，均值±样本标准差基于12个fold×seed结果。")
    add_paragraph(doc, "Backbone 的 All Accuracy 为74.49%，而 M2 历史融合达到83.89%，说明在本协议下历史上下文带来约9.40个百分点的整体增益。该差值用于描述上下文价值，不应被解释为严格的独立模型显著性比较。")

    add_heading(doc, "6. 历史模型主结果", 1)
    add_heading(doc, "6.1 Tier3 总体结果", 2)
    rows = []
    for exp in EXPERIMENTS:
        row = [SHORT[exp]]
        for split in SPLITS:
            row.extend([mean_sd(DATA["aggregates"][exp][split]["tier3_accuracy"]),
                        mean_sd(DATA["aggregates"][exp][split]["tier3_macro_f1"])])
        rows.append(row)
    add_table(doc, ["模型", "N Acc", "N F1", "F Acc", "F F1", "All Acc", "All F1"], rows,
              [2100, 1210, 1210, 1210, 1210, 1210, 1210], font_size=7.8)
    add_caption(doc, "表4  Tier3 性能（%），均值±样本标准差，n=12 fold×seed。粗体未用于标记显著性；最佳均值见正文。")
    add_image(doc, charts["overall"], 6.35, "五种历史模型在Normal、Fault和All上的Tier3 Accuracy分组柱状图")
    add_caption(doc, "图1  Tier3 Accuracy 总体比较。Normal/All 由A1 Once最高，Fault由A1 Every10最高。")
    add_paragraph(doc, "A1-Legacy-Once 在 Normal、Fault、All 的 Accuracy 分别为86.56%、83.32%、85.46%；相比 M2 分别提高1.38、1.68、1.58个百分点。A1 Every10 在 Fault 上达到83.60%，是该 split 最佳，但在 Normal 与 All 上低于 A1 Once。")
    add_paragraph(doc, "DualPos 两组在 Normal 与 All 上均低于 M2。DualPos Every10 的 Fault Accuracy 略高于 M2 0.25个百分点，但其 Fault Macro-F1 仅高0.18个百分点，且跨seed/participant表现并不一致。")

    add_heading(doc, "6.2 Node-level 结果", 2)
    rows = []
    for exp in EXPERIMENTS:
        row = [SHORT[exp]]
        for split in SPLITS:
            row.extend([mean_sd(DATA["aggregates"][exp][split]["node_accuracy"]),
                        mean_sd(DATA["aggregates"][exp][split]["node_macro_f1"])])
        rows.append(row)
    add_table(doc, ["模型", "N Acc", "N F1", "F Acc", "F F1", "All Acc", "All F1"], rows,
              [2100, 1210, 1210, 1210, 1210, 1210, 1210], font_size=7.8)
    add_caption(doc, "表5  Node-level 性能（%），均值±样本标准差，n=12。")
    add_paragraph(doc, "Node-level 与 Tier3 的结论一致：A1 Once 在 All 上达到85.39% Accuracy，比 M2 的83.67%高1.72个百分点；DualPos Once/Every10 分别为81.86%/82.15%。因此主结论不是由Node到Tier3映射方式单独造成。")

    add_heading(doc, "6.3 相对 M2 的配对差值", 2)
    rows = []
    for exp in EXPERIMENTS[1:]:
        z = DATA["deltas_vs_m2"][exp]
        all_z = z["all"]["tier3_accuracy"]
        rows.append([SHORT[exp], f"{z['normal']['tier3_accuracy']['mean']*100:+.2f}",
                     f"{z['fault']['tier3_accuracy']['mean']*100:+.2f}", f"{all_z['mean']*100:+.2f}",
                     f"{all_z['wins']}/{all_z['ties']}/{all_z['losses']}",
                     f"{all_z['fold_wins']}/{all_z['fold_ties']}/{all_z['fold_losses']}"])
    add_table(doc, ["模型", "ΔN Acc", "ΔF Acc", "ΔAll Acc", "All seed W/T/L", "All fold W/T/L"], rows,
              [2300, 1280, 1280, 1280, 1600, 1620], font_size=8.5)
    add_caption(doc, "表6  相对同fold、同seed M2的Tier3 Accuracy差值（百分点）。W/T/L为胜/平/负。")
    add_image(doc, charts["delta"], 6.15, "四种增强模型相对M2的Tier3 Accuracy平均配对差值热力图")
    add_caption(doc, "图2  四种增强模型相对M2的平均配对差值。")
    add_callout(doc, "统计解释", "3个seed共享同一held-out participant和测试样本，不能视为12个完全独立实验单位。真正的跨人独立单位只有4折，因此本报告以描述性均值、配对差值和胜负分布为主，不做夸大的显著性声明。", "gold")

    add_heading(doc, "7. 跨参与者与分阶段分析", 1)
    add_heading(doc, "7.1 Fold heterogeneity", 2)
    add_image(doc, charts["fold"], 6.25, "All split中五种模型按held-out participant的Tier3 Accuracy折线图")
    add_caption(doc, "图3  All split 中每个held-out participant的3-seed平均Tier3 Accuracy。")
    for split in SPLITS:
        rows = []
        for exp in EXPERIMENTS:
            z = DATA["fold_means"][exp][split]
            rows.append([SHORT[exp]] + [f"{z[p]['tier3_accuracy']*100:.2f}" for p in PARTICIPANTS])
        add_caption(doc, f"表{7 + SPLITS.index(split)}  {SPLIT_CN[split]}：按fold的3-seed平均Tier3 Accuracy（%）。")
        add_table(doc, [f"{SPLIT_CN[split]} / 模型", "A", "D", "J", "M"], rows,
                  [2500, 1715, 1715, 1715, 1715], font_size=8.7)
    add_paragraph(doc, "A1 Once 的主要收益来自 A 与 D：All 分别比 M2 高4.80和2.53个百分点；J 与 M 分别低0.72和0.30个百分点。J fold 的 M2 基线已经达到92.67%，增强后略降，呈现明显的高基线/潜在天花板特征。")
    add_paragraph(doc, "A fold 的 M2 All Accuracy 仅74.56%，A1 Once 提升至79.35%；该折训练集过滤最强、保留clips最少。这个现象与“顺序增强在顺序覆盖不足时更有价值”的假设一致，但只有一个最强过滤折，尚不足以建立稳健关联。")

    add_heading(doc, "7.2 Stage-level 结果", 2)
    rows = []
    for exp in EXPERIMENTS:
        z = DATA["stage_aggregates"][exp]
        rows.append([SHORT[exp]] + [mean_sd(z[str(stage)]["tier3_accuracy"]) for stage in (1, 2, 3)])
    add_table(doc, ["模型", "Stage 1 Acc", "Stage 2 Acc", "Stage 3 Acc"], rows,
              [2600, 2253, 2253, 2254], font_size=8.8)
    add_caption(doc, "表10  test_all 的stage-level Tier3 Accuracy（%，12个fold×seed均值±SD）。")
    add_paragraph(doc, "A1 Once 的增益主要集中在 Stage 2：85.87% 对 M2 的83.72%，提高2.15个百分点；Stage 1和3基本持平。Stage 2包含动作密度更高、顺序分支更多的主体操作段，这与graph-valid顺序增强的目标相符。")

    add_heading(doc, "8. 增强行为与顺序覆盖", 1)
    add_heading(doc, "8.1 实际训练视图的改变强度", 2)
    rows = []
    for exp in EXPERIMENTS:
        z = DATA["runtime_augmentation_audits"][exp]
        rows.append([SHORT[exp], pct(z["augmentation_changed_fraction"]["mean"]),
                     f"{z['mean_normalized_kendall_distance']['mean']:.3f}",
                     pct(z["shifted_history_token_fraction"]["mean"]),
                     f"{z['mean_absolute_position_shift']['mean']:.3f}"])
    add_table(doc, ["模型", "Changed histories", "Mean Kendall", "Shifted tokens", "Mean abs shift"], rows,
              [2300, 1800, 1650, 1800, 1810], font_size=8.5)
    add_caption(doc, "表11  训练时augmentation audit的12任务均值（%或表中标注尺度）。")
    add_paragraph(doc, "A1 的约53.43%训练历史发生变化，DualPos约29.73%。A1 的 shifted-token 指标为0并不是没有shuffle：它在重排后重新赋presented位置，因此没有“相对真实位置的位移embedding”；DualPos则显式记录位移，约22.96%的历史token具有非零shift。")

    add_heading(doc, "8.2 Post-hoc 测试前缀覆盖审计", 2)
    rows = []
    for exp in EXPERIMENTS[1:]:
        z = DATA["coverage"][exp]
        rows.append([SHORT[exp], pct(z["mean_changed_fraction"]), pct(z["mean_actual_test_prefix_coverage"]),
                     pct(z["mean_augmented_test_prefix_coverage"]), z["total_new_test_prefixes_covered"]])
    add_table(doc, ["模型", "Changed", "Actual coverage", "Augmented coverage", "新覆盖前缀总数"], rows,
              [2450, 1550, 1800, 1900, 1660], font_size=8.5)
    add_caption(doc, "表12  四折post-hoc coverage均值；最后一列为四折新覆盖测试前缀数之和。")
    add_image(doc, charts["coverage"], 6.25, "四种增强模型的历史改变比例、真实前缀覆盖和增强前缀覆盖柱状图")
    add_caption(doc, "图4  增强改变比例与测试历史前缀覆盖。该审计不参与采样，不构成测试泄漏。")
    add_paragraph(doc, "A1 Every10 将平均 augmented prefix coverage 提高到22.68%，并累计新覆盖34个测试前缀，高于 A1 Once 的19.15%与16个。然而其 All Accuracy 反而低0.89个百分点。这表明覆盖数量不是充分条件：频繁替换可能增加优化噪声，或生成的额外前缀与判别边界并不匹配。")
    add_paragraph(doc, "coverage 的比较单位是“当前clip之前的node-order历史前缀”，而不是一条连贯合成的完整run。Legacy augmenter按样本独立重排历史，因此不能把这些覆盖数字解释为生成了同等数量的完整新流程。")

    add_heading(doc, "9. Once 与 Every10 的直接比较", 1)
    rows = []
    for pair in ("A1 Every10 - Once", "DualPos Every10 - Once"):
        for split in SPLITS:
            z = DATA["refresh_deltas"][pair][split]["tier3_accuracy"]
            rows.append([pair, SPLIT_CN[split], f"{z['mean']*100:+.2f}", f"{z['wins']}/{z['ties']}/{z['losses']}",
                         " ".join(f"{p}:{z['fold_mean_deltas'][p]*100:+.2f}" for p in PARTICIPANTS)])
    add_table(doc, ["比较", "Split", "平均ΔAcc", "Seed W/T/L", "Fold均值差值"], rows,
              [2450, 900, 1350, 1450, 3210], font_size=8.2)
    add_caption(doc, "表13  Every10减Once的Tier3 Accuracy配对差值（百分点）。")
    add_paragraph(doc, "A1 Every10 对 Fault 有+0.28个百分点的轻微平均收益，但 Normal/All 分别为-1.07/-0.89；All 中仅3胜、1平、8负。DualPos Every10 与 Once 的 All 几乎相同（+0.04），主要表现为Fault稍升、Normal稍降。")
    add_paragraph(doc, "因此，本轮没有证据支持“每10轮替换一定优于只shuffle一次”。更稳妥的解释是：A1的固定增强视图更容易优化，而DualPos对刷新频率不敏感。")

    add_heading(doc, "10. 对研究假设的回答", 1)
    add_heading(doc, "10.1 能否证明 graph-valid shuffle 优于 M2？", 2)
    add_callout(doc, "回答", "可以给出有限度、模型特定的支持：A1-Legacy-Once 在严格 sequence-disjoint 条件下取得最高总体均值，并在 Normal、Fault、All 三个split都高于M2；但优势集中于部分参与者，不能写成无条件的普遍优越性。", "green")
    add_paragraph(doc, "最合适的论文式表述是：‘Under exact full-run sequence disjointness, legacy graph-valid shuffle with post-shuffle presented-position encoding improved mean Tier3 accuracy over the real-order M2 baseline by 1.58 percentage points on the combined test split, although gains were participant-dependent.’")
    add_heading(doc, "10.2 哪个模型最适合作为当前主方法？", 2)
    add_paragraph(doc, "当前主方法应选 A1-Legacy-Once。它既实现了训练时新顺序暴露，又避免保留真实recency对原顺序的显式提示；同时在三个split上均给出正向平均增益。A1 Every10 可作为Fault-oriented消融，而不宜替代 Once 作为总体最佳配置。")
    add_heading(doc, "10.3 为什么 DualPos 没有体现优势？", 2)
    add_paragraph(doc, "第一，DualPos只增强active atomic tail，实际改变历史比例约为A1的一半；第二，它保留真实recency，模型仍能读取原始时间关系，shuffle对顺序不变性的训练压力更弱；第三，A1与DualPos同时在增强范围和位置编码上不同，本轮不是干净的单因素比较。因此只能说‘当前DualPos组合无效’，不能断言‘位移编码本身无效’。")

    add_heading(doc, "11. 局限性与后续实验建议", 1)
    add_heading(doc, "11.1 主要局限", 2)
    add_paragraph(doc, "独立跨人单位仅4个；seed重复主要刻画训练随机性，不能替代更多参与者。Fault样本按fold差异明显（62–168），导致Fault均值波动较大。Sequence-disjoint过滤移除了大量训练数据，因此结果同时反映了‘顺序未见’和‘训练样本减少’两个因素。")
    add_paragraph(doc, "此外，测试前缀覆盖是post-hoc诊断指标，并不衡量生成顺序的任务相关性或样本权重；A1与DualPos的比较还混合了active-tail scope和position semantics。")
    add_heading(doc, "11.2 优先建议", 2)
    add_paragraph(doc, "第一优先：补做2×2消融，将shuffle scope（broad vs active-tail-only）与position semantics（presented vs true+shift）正交化。至少加入 broad-shuffle DualPos 与 active-tail presented-position 两个模型。")
    add_paragraph(doc, "第二优先：保留 A1 Once 作为主方法，针对每折统计训练实际顺序数量、增强后unique prefixes、A1增益之间的相关趋势；若参与者增加，再进行以participant为单位的置信区间或层级模型分析。")
    add_paragraph(doc, "第三优先：Every10不再作为默认增强策略。若继续研究刷新频率，建议增加Every20或curriculum式刷新，并记录每次刷新前后的训练损失跳变，以区分覆盖收益与优化扰动。")
    add_paragraph(doc, "第四优先：增加一个‘不做sequence过滤但保持同等训练run数量’的下采样控制，用于拆分顺序隔离效应与数据量下降效应。")

    doc.add_page_break()
    add_heading(doc, "附录A：逐fold、逐seed Tier3 Accuracy", 1)
    records = DATA["records"]
    lookup = {(r["experiment"], r["participant"], r["seed"], r["split"]): r for r in records}
    table_no = 14
    for split in SPLITS:
        rows = []
        for exp in EXPERIMENTS:
            for p in PARTICIPANTS:
                vals = [lookup[(exp, p, seed, split)]["tier3_accuracy"]*100 for seed in SEEDS]
                rows.append([SHORT[exp], p] + [f"{v:.2f}" for v in vals] + [f"{sum(vals)/3:.2f}"])
        add_table(doc, ["模型", "Fold", "Seed 1", "Seed 2", "Seed 42", "Mean"], rows,
                  [2300, 900, 1480, 1480, 1600, 1600], font_size=7.8)
        add_caption(doc, f"表{table_no}  {SPLIT_CN[split]} split逐fold、逐seed Tier3 Accuracy（%）。")
        table_no += 1
        if split != "all":
            doc.add_page_break()

    add_heading(doc, "附录B：指标定义与汇总口径", 1)
    add_paragraph(doc, "Accuracy 为正确预测样本比例；Macro-F1 对当前split中被评估类别的F1做等权平均；Balanced Accuracy 对各类别召回率做等权平均。Node-level针对35个图节点，Tier3-level针对31个动作类别。")
    add_paragraph(doc, "正文的‘均值±SD’以12个fold×seed记录计算，便于与既有实验汇总保持一致；fold表先对3个seed求均值，更适合观察跨人的方向一致性。All并非Normal与Fault指标的简单算术平均，而是在合并样本上重新计算。分析源包括sequence-disjoint与coverage总审计、12组backbone日志、60组history任务及180个split metrics。")

    # Set core properties and save.
    doc.core_properties.title = "Sequence-Disjoint Atomic-Tail 实验分析报告"
    doc.core_properties.subject = "Graph-valid shuffle augmentation under unseen exact run orders"
    doc.core_properties.author = "Objective 3 Experiment Analysis"
    doc.core_properties.keywords = "sequence-disjoint, graph-valid shuffle, atomic-tail, M2, A1, DualPos"
    doc.save(OUT)
    print("Report written successfully.")


if __name__ == "__main__":
    build_report()
