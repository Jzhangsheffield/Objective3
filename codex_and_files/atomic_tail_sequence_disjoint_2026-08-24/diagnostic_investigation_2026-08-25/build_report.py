from __future__ import annotations

import math
from pathlib import Path
from typing import Any, Iterable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUT = HERE / "Graph_Valid_Shuffle_Augmentation_前三步排查报告_2026-08-25.docx"
ASSETS = HERE / "report_assets"
PACKAGE_ROOT = HERE.parent

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "6B7280"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
TEAL = "2A9D8F"
RED = "9B1C1C"
GOLD = "7A5A00"
WHITE = "FFFFFF"
TABLE_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run: Any, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table: Any, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != TABLE_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_DXA))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table: Any, color: str = "D7DEE8", size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_paragraph_border_bottom(paragraph: Any, color: str = BLUE, size: int = 16, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=GRAY)
    for element in (
        ("w:fldChar", "begin", None),
        ("w:instrText", None, " PAGE "),
        ("w:fldChar", "separate", None),
        ("w:t", None, "1"),
        ("w:fldChar", "end", None),
    ):
        node = OxmlElement(element[0])
        if element[1] is not None:
            node.set(qn("w:fldCharType"), element[1])
        if element[0] == "w:instrText":
            node.set(qn("xml:space"), "preserve")
        if element[2] is not None:
            node.text = element[2]
        field_run = OxmlElement("w:r")
        field_run.append(node)
        paragraph._p.append(field_run)
    run = paragraph.add_run(" 页")
    set_run_font(run, 9, color=GRAY)


def configure_styles(doc: Document) -> None:
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    p = header.paragraphs[0]
    p.text = "SEQUENCE-DISJOINT DIAGNOSTIC  |  GRAPH-VALID SHUFFLE"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_run_font(run, 8.5, bold=True, color=GRAY)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("DIAGNOSTIC REPORT  /  2026-08-25")
    set_run_font(run, 10, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Graph-Valid Shuffle Augmentation\n前三步排查报告")
    set_run_font(run, 27, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(36)
    run = subtitle.add_run("顺序敏感性 · 局部新颖性分组 · 增强历史真实性与多样性")
    set_run_font(run, 14, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(28)
    set_paragraph_border_bottom(rule, BLUE, 18, 10)

    metadata = [
        ("实验包", "atomic_tail_sequence_disjoint_2026-08-24"),
        ("模型", "M2-Direct-RealOrder / A1-Legacy-Once / A3-DualPos-Once"),
        ("评估", "4 LOSO folds × 3 seeds；Normal / Fault / All"),
        ("目的", "解释 graph-valid shuffle augmentation 作用不显著的直接原因"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_geometry(table, [1900, 7460])
    set_table_borders(table, color="E4E8EE", size=4)
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_fill(row.cells[0], LIGHT_GRAY)
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0 = p0.add_run(label)
        r1 = p1.add_run(value)
        set_run_font(r0, 10, bold=True, color=DARK_BLUE)
        set_run_font(r1, 10, color=NAVY)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("结论先行：模型会响应强烈的无效换序，但对现有 graph-valid 换序几乎不改变预测。")
    set_run_font(run, 11.5, bold=True, color=RED)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> Any:
    return doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold_lead: str | None = None, color: str | None = None, italic: bool = False) -> Any:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, 11, bold=True, color=color or NAVY)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, 11, color=color or None, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 11, color=color or None, italic=italic)
    return paragraph


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, 11)


def add_callout(doc: Document, title: str, text: str, tone: str = "info") -> None:
    color = {"info": DARK_BLUE, "positive": TEAL, "caution": GOLD, "risk": RED}[tone]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_DXA])
    set_table_borders(table, color=BLUE_GRAY, size=6)
    set_cell_fill(table.cell(0, 0), CALLOUT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title + "  ")
    set_run_font(run, 11, bold=True, color=color)
    run = p.add_run(text)
    set_run_font(run, 10.5, color=NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[Any]],
    widths: list[int],
    *,
    font_size: float = 8.5,
    first_col_bold: bool = False,
    caption: str | None = None,
) -> Any:
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(caption)
        set_run_font(run, 9.5, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(header))
        set_run_font(run, font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, font_size, bold=bool(first_col_bold and index == 0), color=NAVY)
            if row_index % 2 == 1:
                set_cell_fill(cells[index], "FAFBFC")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def set_image_alt(inline_shape: Any, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description[:120])


def add_figure(doc: Document, filename: str, caption: str, alt: str, width: float = 6.45) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(ASSETS / filename), width=Inches(width))
    set_image_alt(shape, alt)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_font(run, 9, italic=True, color=GRAY)


def pct(value: float, digits: int = 2) -> str:
    return f"{100 * float(value):.{digits}f}%"


def pp(value: float, digits: int = 2) -> str:
    return f"{100 * float(value):+.{digits}f} pp"


def mean_sd(mean: float, sd: float, digits: int = 2) -> str:
    return f"{100*float(mean):.{digits}f} ± {100*float(sd):.{digits}f}%"


def load_data() -> dict[str, pd.DataFrame]:
    return {
        "order": pd.read_csv(HERE / "order_sensitivity_summary.csv"),
        "validation": pd.read_csv(HERE / "order_sensitivity_reproduction_check.csv"),
        "groups": pd.read_csv(HERE / "grouped_performance_summary.csv"),
        "deltas": pd.read_csv(HERE / "grouped_performance_model_deltas_summary.csv"),
        "metadata": pd.read_csv(HERE / "test_group_metadata.csv"),
        "aug": pd.read_csv(HERE / "augmentation_history_audit_summary.csv"),
        "aug_detail": pd.read_csv(HERE / "augmentation_history_audit_detailed.csv"),
        "diversity": pd.read_csv(HERE / "augmentation_diversity_by_sample.csv"),
        "manual": pd.read_csv(HERE / "manual_audit_candidates.csv"),
    }


def add_executive_summary(doc: Document, data: dict[str, pd.DataFrame]) -> None:
    add_heading(doc, "执行摘要", 1)
    add_callout(
        doc,
        "核心判断",
        "现有 graph-valid shuffle 并不是完全无效，但它没有形成足够强、足够真实、且模型无法轻易忽略的顺序干预。因此 A1 只获得小幅且不稳定的平均收益，DualPos 则几乎恢复了真实顺序，削弱了增强目标。",
        "caution",
    )
    add_para(doc, "三步排查形成了相互一致的证据链：")
    add_bullets(doc, [
        "模型确实使用历史内容：移除历史后 Tier3 Accuracy 下降约 11.6–15.0 个百分点；但对真正的 graph-valid 换序，Tier3 top-1 仅改变 0.10–2.31%。",
        "当历史特征与其真实位置编码一起移动时，三个模型的输出完全不变（top-1 改变率 0，JS divergence 约 10⁻¹⁰），这与单查询 attention 的集合聚合结构一致。",
        "四折 1,895 个唯一测试 clip 中，76.6% 的 Local-3 历史已经在训练折出现，真正 Local-3 unseen 只有 7.1%；full-run disjoint 并未制造大规模局部顺序新颖性。",
        "A1 相对 M2 的 All 增益集中在 Local-3 seen（+1.82 pp）、Stage 2（+2.15 pp）和 history length 3–5（+4.36 pp）；Local-3 unseen 反而为 −1.54 pp。",
        "增强视图全部未新增 hard graph、stage 或 immediate-predecessor 违规，但真正换序的视图中 22.5–24.2% 相邻转移未在训练折出现；A1 仅 3.5%、DualPos 约 0.04% 与某个真实训练 prefix 完全相同。",
        "Every10 在可增强样本上已经产生较高多样性（A1 平均 12.61/15、DualPos 14.65/15 个唯一顺序），因此其性能不优于 Once 不能简单归因于‘没生成新顺序’，更可能与增强质量、分布偏移和持续替换带来的优化噪声有关。",
    ])
    add_para(doc, "因此，本次排查最支持的解释不是“模型完全不看历史”，而是：")
    add_callout(
        doc,
        "最可能机制",
        "历史内容很重要；合法顺序的可变空间较小，Direct Fusion 对物理排列本身又是置换不变的；DualPos 进一步保留真实 recency。与此同时，现有图对 Stage 1 并行动作的约束较宽，生成了不少数学合法但经验稀有的邻接组合。",
        "info",
    )
    add_heading(doc, "排查范围与数据完整性", 2)
    add_table(
        doc,
        ["项目", "配置 / 数量", "用途"],
        [
            ["Folds", "A / D / J / M held out", "LOSO；每折训练 run 与测试 run 顺序隔离"],
            ["Seeds", "1 / 2 / 42", "每个模型 12 个独立训练 job"],
            ["测试 clip", "1,895 unique；5,685 fold×seed", "Normal 1,441；Fault 454；All 1,895"],
            ["顺序敏感性", "36 checkpoints × 6 counterfactual views", "M2 / A1 Once / DualPos Once"],
            ["增强审计", "103,356 generated views", "4 个增强配置；Once 与 Every10"],
            ["人工审查", "31 分层候选", "动作标签层面语义检查"],
        ],
        [1800, 3000, 4560],
        font_size=9,
        first_col_bold=True,
        caption="表 1  排查数据范围",
    )
    validation = data["validation"]
    add_callout(
        doc,
        "复现门槛已通过",
        f"NumPy 反事实推理对 36 个 checkpoint 的原始 actual-order top-1 预测复现率均为 100%，共 {int(validation.samples.sum()):,} 条 checkpoint-sample 记录、0 mismatch。",
        "positive",
    )
def add_method_one(doc: Document, data: dict[str, pd.DataFrame]) -> None:
    order = data["order"]
    add_heading(doc, "1. 排查一：测试时顺序敏感性分析", 1)
    add_heading(doc, "1.1 目的与判别逻辑", 2)
    add_para(doc, "目的不是再次比较训练精度，而是在固定 checkpoint、固定当前 RGB 特征和固定历史动作集合的情况下，只改变历史动作顺序，观察模型输出是否变化。")
    add_bullets(doc, [
        "若 graph-valid 换序后 logits / probabilities 几乎不变，则训练时加入此增强很难产生明显收益。",
        "若随机或反转顺序会大幅改变输出、而 graph-valid 不会，说明模型能读取位置—动作对应关系，但当前合法换序的有效扰动太弱或语义差异太小。",
        "若保留每个 token 的真实位置编码时仅改变数组排列，输出应当不变；该项是架构置换不变性的负对照。",
    ])
    add_heading(doc, "1.2 反事实视图的构造", 2)
    add_table(
        doc,
        ["视图", "历史动作集合", "位置编码", "含义"],
        [
            ["Actual", "不变", "真实呈现顺序", "原始测试基准"],
            ["Graph-valid true-pos control", "不变，仅重排存储槽", "每个 token 保留原真实 recency", "验证物理排列是否被模型读取"],
            ["Graph-valid native", "图合法换序", "M2/A1: presented；DualPos: true recency + shift", "与各方法训练语义一致的核心测试"],
            ["Random presented", "随机置换", "按新呈现顺序重编", "强度较高、通常非 graph-valid 的正对照"],
            ["Reverse presented", "完全反转", "按反转顺序重编", "最强顺序破坏正对照"],
            ["No history", "删除全部历史", "无", "校准历史内容本身的重要性"],
        ],
        [2200, 1900, 2350, 2910],
        font_size=8.5,
        first_col_bold=True,
        caption="表 2  测试时反事实顺序视图",
    )
    add_heading(doc, "1.3 统计指标与口径", 2)
    add_bullets(doc, [
        "Tier3 / Node top-1 改变率：反事实预测类别与 actual-order 预测不同的比例。",
        "Jensen–Shannon divergence：两个 node 概率分布的对称差异；0 表示完全一致。",
        "Total variation：0.5 × 概率绝对差之和；可直观理解为概率质量移动量。",
        "Accuracy delta：同一批样本上反事实 Tier3 Accuracy 减去 actual-order Accuracy。",
        "Graph-valid changed-only：只统计本次随机拓扑排序确实改变了历史顺序的样本，避免 unchanged 样本稀释敏感性。",
    ])
    add_heading(doc, "1.4 结果", 2)
    add_figure(
        doc,
        "01_order_sensitivity.png",
        "图 1  graph-valid 原生换序几乎不改变 Tier3 top-1，而随机/反转换序会产生明显变化。",
        "三种模型在 graph-valid、随机和反转顺序下的 Tier3 top-1 改变率柱状图。",
    )
    changed = order[(order.analysis_scope == "graph_permutation_changed_only") & (order.variant == "graph_valid_native")]
    random_rows = order[(order.analysis_scope == "eligible_history_len_ge_2") & (order.variant == "random_presented")]
    reverse_rows = order[(order.analysis_scope == "eligible_history_len_ge_2") & (order.variant == "reverse_presented")]
    no_hist = order[(order.analysis_scope == "eligible_history_len_ge_2") & (order.variant == "no_history")]
    model_labels = {
        "M2-Direct-RealOrder": "M2",
        "A1-Legacy-Once": "A1 Once",
        "A3-DualPos-Once": "DualPos Once",
    }
    rows = []
    for model, label in model_labels.items():
        c = changed[changed.model == model].iloc[0]
        r = random_rows[random_rows.model == model].iloc[0]
        v = reverse_rows[reverse_rows.model == model].iloc[0]
        n = no_hist[no_hist.model == model].iloc[0]
        rows.append([
            label,
            int(c.samples),
            pct(c.tier3_top1_change_rate),
            f"{c.mean_js:.6f}",
            pp(c.tier3_accuracy_delta_vs_actual),
            pct(r.tier3_top1_change_rate),
            pp(r.tier3_accuracy_delta_vs_actual),
            pct(v.tier3_top1_change_rate),
            pp(n.tier3_accuracy_delta_vs_actual),
        ])
    add_table(
        doc,
        ["模型", "合法换序 n", "合法 Tier3\n改变", "合法 JS", "合法 Acc Δ", "随机改变", "随机 Acc Δ", "反转改变", "无历史 Acc Δ"],
        rows,
        [1050, 850, 1100, 900, 1000, 1050, 1000, 1050, 1360],
        font_size=7.8,
        first_col_bold=True,
        caption="表 3  顺序敏感性核心结果",
    )
    add_heading(doc, "1.5 对 true-position 负对照的解释", 2)
    add_para(doc, "三个模型在 `graph_valid_truepos_control` 上的 Node 与 Tier3 top-1 改变率均为 0，平均 JS 约为 10⁻¹⁰。这个结果不是偶然：当前模型使用单个 current query 对所有 history key/value 做 MultiheadAttention，没有 history-to-history self-attention。只要每个动作特征与其位置向量作为一对共同移动，attention 对 token 的物理排列是置换不变的。")
    add_callout(
        doc,
        "架构结论",
        "模型读取的是‘动作特征 + 位置/shift 编码’这个集合，而不是数组槽位的先后。A1 只有在 shuffle 后重新分配 presented position，才真正改变动作与位置的配对；DualPos 又把 true recency 保留下来，因此合法换序对输出的影响更小。",
        "info",
    )
    add_heading(doc, "1.6 对 graph-valid 与强扰动差异的解释", 2)
    add_bullets(doc, [
        "M2 与 A1 在实际发生 graph-valid 换序的样本上，Tier3 top-1 只改变 1.94% 和 2.31%；DualPos 仅 0.10%。",
        "随机/反转换序会改变 14.62–20.64% 的 Tier3 top-1，并使 Accuracy 下降 10.24–16.12 pp，说明模型并非完全忽略位置。",
        "因此问题更接近‘合法换序没有制造模型所关心的差异’，而不是‘模型没有任何顺序通道’。",
        "No-history 造成 11.62–14.98 pp 的准确率损失，进一步说明历史内容有明显价值；但历史内容有用不等同于精细顺序有用。",
    ])
def group_row(groups: pd.DataFrame, model: str, condition: str, grouping: str, group: str) -> pd.Series:
    values = groups[(groups.model == model) & (groups.condition == condition) & (groups.grouping == grouping) & (groups.group.astype(str) == str(group))]
    if values.empty:
        raise KeyError((model, condition, grouping, group))
    return values.iloc[0]


def delta_row(deltas: pd.DataFrame, model: str, condition: str, grouping: str, group: str) -> pd.Series:
    values = deltas[(deltas.comparison_model == model) & (deltas.condition == condition) & (deltas.grouping == grouping) & (deltas.group.astype(str) == str(group))]
    if values.empty:
        raise KeyError((model, condition, grouping, group))
    return values.iloc[0]


def add_method_two(doc: Document, data: dict[str, pd.DataFrame]) -> None:
    groups, deltas, metadata = data["groups"], data["deltas"], data["metadata"]
    add_heading(doc, "2. 排查二：按局部新颖性、历史长度、Stage 与 Active Tail 重分组", 1)
    add_heading(doc, "2.1 为什么要重新分组", 2)
    add_para(doc, "Sequence-disjoint 只保证完整训练 run 的 node 顺序不与测试 run 完全相同，并不保证当前 clip 之前的短历史、局部转移或 history set 未见。若模型主要依赖最后 1–3 个动作或 Stage 信息，那么全 run 隔离仍可能对模型很容易。")
    add_heading(doc, "2.2 分组定义", 2)
    add_table(
        doc,
        ["分组", "定义", "要回答的问题"],
        [
            ["Exact full prefix", "当前动作之前的完整 node history 是否等于训练折任一 clip 的完整 history prefix", "完整局部上下文是否见过"],
            ["Local-1/2/3", "当前动作前紧邻的最后 k 个历史节点是否作为同样的局部尾部出现在训练折", "模型最可能依赖的局部顺序是否新颖"],
            ["History length", "0 / 1–2 / 3–5 / 6–10 / 11–20 / 21+", "增强在哪种上下文长度有效"],
            ["Stage", "当前动作的 Stage 1 / 2 / 3", "增益是否只来自特定工艺阶段"],
            ["Active tail", "不观察 current target，仅由历史末端检测未完成 atomic prefix", "atomic-tail gating 对应的样本是否获益"],
        ],
        [1800, 4380, 3180],
        font_size=8.5,
        first_col_bold=True,
        caption="表 4  测试样本重分组口径",
    )
    add_para(doc, "说明：本报告中的 Local-3 指 history prefix 的‘局部尾部’，即当前动作之前紧邻的最后 3 个节点；history length < 3 的样本单列为 insufficient，避免把它们误判为 unseen。")
    add_heading(doc, "2.3 测试集局部覆盖", 2)
    add_figure(
        doc,
        "05_test_composition.png",
        "图 2  Full-run sequence-disjoint 下，局部历史仍有很高训练覆盖。",
        "测试样本 exact full history prefix、Local-3 prefix 与 active tail 构成的堆叠条形图。",
    )
    total = len(metadata)
    add_bullets(doc, [
        f"Exact full history prefix：seen {int(metadata.exact_full_prefix_seen.sum())}/{total}（{metadata.exact_full_prefix_seen.mean():.1%}），unseen {int((~metadata.exact_full_prefix_seen).sum())}/{total}。",
        f"Local-3：seen {(metadata.local_prefix3_status == 'seen').sum()}/{total}（{(metadata.local_prefix3_status == 'seen').mean():.1%}），unseen {(metadata.local_prefix3_status == 'unseen').sum()}/{total}（{(metadata.local_prefix3_status == 'unseen').mean():.1%}），length < 3 为 {(metadata.local_prefix3_status == 'insufficient').sum()}。",
        f"Active tail 样本 {int(metadata.active_tail.sum())}/{total}（{metadata.active_tail.mean():.1%}）；其余 30.6% 不满足 active incomplete atomic prefix。",
    ])
    add_callout(
        doc,
        "关键含义",
        "测试 run 的完整顺序虽然被隔离，但模型真正需要的三步局部上下文大多仍然见过。只有 134 个唯一测试 clip 属于 Local-3 unseen，其中 Fault 仅 30 个，因此对 unseen-order 泛化的统计功效有限。",
        "caution",
    )
    add_heading(doc, "2.4 Local-3：Normal / Fault / All", 2)
    rows = []
    for condition in ("normal", "fault", "all"):
        for status in ("seen", "unseen", "insufficient"):
            m2 = group_row(groups, "M2-Direct-RealOrder", condition, "local_prefix_3", status)
            a1 = group_row(groups, "A1-Legacy-Once", condition, "local_prefix_3", status)
            dual = group_row(groups, "A3-DualPos-Once", condition, "local_prefix_3", status)
            delta = delta_row(deltas, "A1-Legacy-Once", condition, "local_prefix_3", status)
            rows.append([
                condition.title(), status,
                int(round(m2.total_samples / 3)),
                mean_sd(m2.tier3_accuracy_mean, m2.tier3_accuracy_sd),
                mean_sd(a1.tier3_accuracy_mean, a1.tier3_accuracy_sd),
                pp(delta.tier3_accuracy_delta_mean),
                mean_sd(dual.tier3_accuracy_mean, dual.tier3_accuracy_sd),
            ])
    add_table(
        doc,
        ["Split", "Local-3", "Unique n", "M2 Acc", "A1 Acc", "A1−M2", "DualPos Acc"],
        rows,
        [900, 1200, 850, 1700, 1700, 1250, 1760],
        font_size=8,
        first_col_bold=True,
        caption="表 5  Local-3 分组 Tier3 Accuracy（12 fold×seed，mean ± SD）",
    )
    add_bullets(doc, [
        "All：A1 在 Local-3 seen 上 +1.82 pp，但 unseen 为 −1.54 pp；这与‘增强主要改善未见顺序泛化’的预期不一致。",
        "Normal unseen：A1 +2.08 pp；Fault unseen：A1 −9.38 pp。Fault unseen 只有 30 个唯一 clip，且跨 fold 方差很大，因此应视为风险信号而非稳定效应。",
        "DualPos 在 seen 和 unseen 上都低于 M2，未显示 true-recency + displacement 对局部新颖性有帮助。",
    ])
    add_heading(doc, "2.5 Exact prefix、History length、Stage 与 Active tail", 2)
    add_figure(
        doc,
        "02_group_deltas.png",
        "图 3  A1 的收益主要来自 Local-3 seen、短中等历史、Stage 2 和 active-tail 样本。",
        "A1-Legacy-Once 相对 M2 在不同测试分组上的 Tier3 Accuracy 平均差横向条形图。",
    )
    selected = [
        ("exact_full_prefix", "False", "Exact prefix unseen"),
        ("exact_full_prefix", "True", "Exact prefix seen"),
        ("history_length", "0", "History 0"),
        ("history_length", "1-2", "History 1–2"),
        ("history_length", "3-5", "History 3–5"),
        ("history_length", "6-10", "History 6–10"),
        ("history_length", "11-20", "History 11–20"),
        ("history_length", "21+", "History 21+"),
        ("stage", "1", "Stage 1"),
        ("stage", "2", "Stage 2"),
        ("stage", "3", "Stage 3"),
        ("active_tail", "True", "Active tail"),
        ("active_tail", "False", "No active tail"),
    ]
    rows = []
    for grouping, group, label in selected:
        m2 = group_row(groups, "M2-Direct-RealOrder", "all", grouping, group)
        a1 = group_row(groups, "A1-Legacy-Once", "all", grouping, group)
        dual = group_row(groups, "A3-DualPos-Once", "all", grouping, group)
        a1_delta = delta_row(deltas, "A1-Legacy-Once", "all", grouping, group)
        dual_delta = delta_row(deltas, "A3-DualPos-Once", "all", grouping, group)
        rows.append([
            label,
            int(round(m2.total_samples / 3)),
            pct(m2.tier3_accuracy_mean),
            pct(a1.tier3_accuracy_mean),
            pp(a1_delta.tier3_accuracy_delta_mean),
            pct(dual.tier3_accuracy_mean),
            pp(dual_delta.tier3_accuracy_delta_mean),
        ])
    add_table(
        doc,
        ["分组", "Unique n", "M2", "A1", "A1−M2", "DualPos", "Dual−M2"],
        rows,
        [2200, 850, 1150, 1150, 1250, 1300, 1460],
        font_size=8,
        first_col_bold=True,
        caption="表 6  All split 分组 Tier3 Accuracy 与相对 M2 差值",
    )
    add_heading(doc, "2.6 分组排查结论", 2)
    add_bullets(doc, [
        "A1 的 +1.58 pp 总体收益并未落在 Local-3 unseen；相反，Exact prefix seen 的增益（+2.36 pp）显著高于 unseen（+0.19 pp）。",
        "最大的历史长度增益位于 3–5（+4.36 pp）和 6–10（+2.92 pp），到 11–20 变为 −0.88 pp，提示长历史重排可能引入噪声。",
        "Stage 2 +2.15 pp，而 Stage 1 / Stage 3 仅 +0.05 / +0.30 pp。整体平均增益主要由 Stage 2 样本量和表现驱动。",
        "Active tail +2.10 pp，高于 no-active-tail +0.43 pp；但这不能直接证明 tail 顺序学习有效，因为 A1 一旦启用会重排 tail 之前的 broad prefix。",
        "History length 0 仍存在模型间差异，说明分组比较是训练结果的描述性关联，不是单独的因果消融；真正因果判断仍需同一 checkpoint 的反事实输入或严格 2×2 训练消融。",
    ])
def add_method_three(doc: Document, data: dict[str, pd.DataFrame]) -> None:
    aug_detail, diversity = data["aug_detail"], data["diversity"]
    add_heading(doc, "3. 排查三：增强历史的真实性、扰动强度与重复度", 1)
    add_heading(doc, "3.1 如何重建“现有增强历史”", 2)
    add_para(doc, "本次没有从训练日志推测增强，而是按原训练代码和 stable seed 重新生成每个训练样本的确定性增强视图：A1 使用 broad graph-valid shuffle + presented position；DualPos 仅在 active tail 条件满足时启用，保留 true recency 并写入 displacement。Once 对每个 seed 生成 1 个 refresh round；Every10 在 50 epochs 下生成 5 个 refresh rounds。")
    add_table(
        doc,
        ["配置", "Views / sample", "总生成数", "真正换序数", "换序率"],
        [],
        [2500, 1600, 1700, 1700, 1860],
        font_size=8.5,
        caption="表 7  增强视图重建规模",
    )
    # Fill the last-created table after construction to keep exact geometry helper centralized.
    table = doc.tables[-1]
    for model, views in (("A1 Once", 3), ("A1 Every10", 15), ("DualPos Once", 3), ("DualPos Every10", 15)):
        source = {
            "A1 Once": "A1-Legacy-Once",
            "A1 Every10": "A1-Legacy-Every10-Replace",
            "DualPos Once": "A3-DualPos-Once",
            "DualPos Every10": "A3-DualPos-Every10",
        }[model]
        values = aug_detail[aug_detail.model == source]
        cells = table.add_row().cells
        row_values = [model, views, f"{len(values):,}", f"{int(values.changed.sum()):,}", pct(values.changed.mean())]
        for index, value in enumerate(row_values):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, 8.5, bold=index == 0, color=NAVY)
    set_table_geometry(table, [2500, 1600, 1700, 1700, 1860])

    add_heading(doc, "3.2 ‘真实性’的三层操作化定义", 2)
    add_bullets(doc, [
        "硬结构合法性：增强 history 是否满足 feature-history graph；追加 current target 后是否仍合法；是否新增 stage 倒退、atomic/immediate-predecessor 破坏。",
        "经验真实性：增强相邻转移是否在同一过滤训练折出现；一阶 Laplace 转移平均 log-probability 是否下降；完整增强 history 是否等于某个真实训练 prefix。",
        "人工语义真实性：分层抽取最大扰动、经验异常、经验支持与 unchanged 候选，直接阅读动作标签、current target 和 before/after 顺序。",
    ])
    add_callout(
        doc,
        "边界",
        "硬图合法只能证明没有违反当前编码的依赖。它不能证明工具位置、样品状态、故障状态或设备状态在物理上可达；经验真实性也只是训练数据支持度，不等于工艺专家认证。",
        "caution",
    )
    add_heading(doc, "3.3 硬结构合法性：增强器没有新增已编码违规", 2)
    changed = aug_detail[aug_detail.changed]
    rows = []
    model_map = {
        "A1-Legacy-Once": "A1 Once",
        "A1-Legacy-Every10-Replace": "A1 Every10",
        "A3-DualPos-Once": "DualPos Once",
        "A3-DualPos-Every10": "DualPos Every10",
    }
    for source, label in model_map.items():
        values = changed[changed.model == source]
        rows.append([
            label,
            pct(values.graph_valid_regressed.mean()),
            pct(values.target_graph_valid_regressed.mean()),
            pct((values.stage_backward_increase > 0).mean()),
            pct((values.immediate_constraint_violation_increase > 0).mean()),
            pct(values.target_immediate_predecessor_regressed.mean()),
        ])
    add_table(
        doc,
        ["配置", "History graph\n新增违规", "追加 target\n新增违规", "Stage 倒退\n新增", "History immediate\n新增违规", "Target immediate\n回归"],
        rows,
        [2100, 1450, 1450, 1400, 1500, 1460],
        font_size=8,
        first_col_bold=True,
        caption="表 8  真正换序视图的硬约束回归率",
    )
    add_para(doc, "所有列均为 0%。这说明当前 sampler 在已编码约束上是正确的：本次问题不是实现错误导致 graph-invalid history。")
    add_heading(doc, "3.4 经验真实性：约四分之一新邻接未在训练折出现", 2)
    add_figure(
        doc,
        "03_augmentation_quality.png",
        "图 4  真正换序的增强历史中，22.5–24.2% 相邻转移未在训练折出现。",
        "四种增强配置的 novel transition fraction 与真实训练 prefix 匹配率柱状图。",
    )
    rows = []
    for source, label in model_map.items():
        values = changed[changed.model == source]
        rows.append([
            label,
            pct(values.novel_transition_fraction.mean(), 1),
            f"{values.actual_mean_transition_log_probability.mean():.3f}",
            f"{values.augmented_mean_transition_log_probability.mean():.3f}",
            f"{values.augmented_mean_transition_log_probability.mean() - values.actual_mean_transition_log_probability.mean():+.3f}",
            pct(values.augmented_signature_seen_as_real_train_prefix.mean(), 2),
        ])
    add_table(
        doc,
        ["配置", "Novel transitions", "Actual logP/edge", "Aug logP/edge", "Δ logP/edge", "Real-prefix match"],
        rows,
        [2100, 1500, 1550, 1500, 1350, 1360],
        font_size=8.3,
        first_col_bold=True,
        caption="表 9  真正换序视图的经验真实性指标",
    )
    add_bullets(doc, [
        "A1 changed views：22.6% 相邻转移未见，平均 logP/edge 从 −1.290 降至 −1.920；只有 3.51% 与真实训练 prefix 完全相同。",
        "DualPos changed views：24.2% 相邻转移未见，logP/edge 从 −1.245 降至 −1.930；real-prefix match 约 0.04%。",
        "Every10 没有改善这些质量指标：A1 与 DualPos 的 novel-transition 和 logP 与 Once 基本相同，只是生成了更多样本。",
        "Real-prefix match 低并不自动意味着错误，因为方法目的就是生成新顺序；但它说明增强分布主要依赖图的外推，而不是训练数据直接支持。",
    ])
    add_heading(doc, "3.5 扰动强度", 2)
    rows = []
    for source, label in model_map.items():
        all_values = aug_detail[aug_detail.model == source]
        values = all_values[all_values.changed]
        rows.append([
            label,
            pct(all_values.changed.mean()),
            f"{values.kendall_distance.mean():.3f}",
            pct(values.changed_position_fraction.mean(), 1),
            f"{values.mean_absolute_index_shift.mean():.2f}",
            pct(values.latest_history_token_moved.mean(), 1),
        ])
    add_table(
        doc,
        ["配置", "全样本换序率", "Kendall\n(changed)", "位置改变比例", "平均索引位移", "最近 token 被移动"],
        rows,
        [2100, 1600, 1400, 1450, 1400, 1410],
        font_size=8.3,
        first_col_bold=True,
        caption="表 10  增强扰动强度",
    )
    add_bullets(doc, [
        "全训练样本上 A1 约 53.2% 发生换序，DualPos 约 29.5%；这解释了 DualPos 的有效增强暴露更少。",
        "在确实换序的样本中，A1 / DualPos 分别有约 57.4% / 54.0% 的历史位置改变，平均索引移动约 1.74 / 1.88；因此 changed view 并非只有极小邻接交换。",
        "A1 中约 32.6% changed view 移动了最近历史 token；DualPos 为 0%，因为 active atomic tail 被固定在末端。",
        "虽然 DualPos changed view 的动作排列变化不小，但模型测试输出变化仅 0.10% Tier3 top-1，说明 true recency + shift 语义使模型几乎恢复或忽略了增强呈现顺序。",
    ])
    add_heading(doc, "3.6 重复度与有效多样性", 2)
    add_figure(
        doc,
        "04_diversity.png",
        "图 5  Every10 提高了绝对唯一顺序数，但大量不可增强样本使总体重复率较高。",
        "四种增强配置每个样本的平均唯一历史数及总体重复比例柱状图。",
    )
    rows = []
    for source, label in model_map.items():
        values = diversity[diversity.model == source]
        eligible = values[values.any_changed]
        rows.append([
            label,
            f"{values.unique_sequences.mean():.2f}/{values.views.mean():.0f}",
            pct(values.duplicate_ratio.mean(), 1),
            pct(values.all_unchanged.mean(), 1),
            f"{eligible.unique_sequences.mean():.2f}/{eligible.views.mean():.0f}",
            pct(eligible.duplicate_ratio.mean(), 1),
        ])
    add_table(
        doc,
        ["配置", "全样本唯一/视图", "全样本重复", "始终 unchanged", "可增强样本唯一/视图", "可增强重复"],
        rows,
        [2000, 1650, 1400, 1450, 1660, 1200],
        font_size=8.2,
        first_col_bold=True,
        caption="表 11  每个训练样本跨 seed / refresh 的多样性",
    )
    add_para(doc, "这组结果改变了对 Every10 的解释：在能发生换序的样本上，A1 Every10 平均 12.61/15 个唯一顺序，DualPos Every10 为 14.65/15，重复率仅 16.0% 和 2.3%。因此 Every10 并不缺新顺序；其总体重复主要来自 43.4%（A1）或 70.2%（DualPos）的样本始终 unchanged。")
    add_callout(
        doc,
        "Every10 诊断",
        "刷新确实增加了可增强样本的顺序覆盖，但没有提高经验真实性，也没有带来更好 All 性能。更合理的原因是：新视图质量与真实分布不完全一致，并且每 10 epochs 替换训练视图会持续改变优化目标。",
        "caution",
    )
    add_heading(doc, "3.7 人工语义审查", 2)
    add_para(doc, "按每 fold × 模型分层抽取 31 个候选：最大 Kendall、novel-transition 最高、经验支持较好和 unchanged/fallback。人工阅读 before/after 动作标签与 current target 后，得到以下分类：")
    add_table(
        doc,
        ["判断", "数量", "典型特征"],
        [
            ["经验支持且可接受", 7, "新相邻转移均见；atomic / immediate 片段完整；无明显语义矛盾"],
            ["形式合法但真实性存疑", 8, "Stage 1 准备动作远距离重排；图允许但实际设备状态未验证"],
            ["经验支持较弱", 8, "50–100% 新相邻转移未见；短历史中尤其明显"],
            ["未形成新样本", 8, "重复节点触发 fallback；增强等于真实历史"],
        ],
        [2600, 1000, 5760],
        font_size=8.8,
        first_col_bold=True,
        caption="表 12  31 个候选的人工语义审查结果",
    )
    add_bullets(doc, [
        "代表性低支持案例：`main switch → water pump → air compressor` 被改为 `water pump → main switch → air compressor`，两个新邻接在该训练折均未出现，但现有 graph 仍判定合法。",
        "代表性大扰动案例：多个 startup 动作近似反转，unlock/put-lock 原子片段仍保持；说明 hard constraints 正确，但 graph 对并行动作的自由度较大。",
        "代表性可信案例：Stage 1 可交换动作调整位置，同时完整保留 Stage 2 线性链和 current target immediate predecessor，且所有新邻接都在训练折出现。",
        "DualPos 的 active-tail gating 保护 tail 内部顺序，但一旦启用，仍会重排 tail 之前的整个 remaining prefix；它不是只对 tail 做局部小交换。",
    ])
    add_para(doc, "更完整的 31 个候选与代表性 before/after 序列记录在同目录 `manual_semantic_review.md` 和 `manual_audit_candidates.csv`。", italic=True, color=GRAY)
def add_integrated_diagnosis(doc: Document) -> None:
    add_heading(doc, "4. 综合诊断：为什么 graph-valid shuffle 的作用不显著", 1)
    add_heading(doc, "4.1 证据链", 2)
    add_table(
        doc,
        ["观察", "直接证据", "对原因的指向"],
        [
            ["历史有用，但合法顺序变化弱", "No-history −11.6 至 −15.0 pp；合法换序 top-1 改变 0.10–2.31%", "模型依赖历史内容/阶段，不强依赖当前合法顺序细节"],
            ["架构天然集合化", "token + true position 一起移动时输出完全不变", "物理排列本身不被读取；只有 feature-position 重新配对才有效"],
            ["DualPos 几乎恢复真实顺序", "合法 changed-only Tier3 改变 0.10%", "true recency 让模型绕过 shuffle；shift 分支未形成强干预"],
            ["局部新颖性不足", "Local-3 seen 76.6%，unseen 仅 7.1%", "sequence-disjoint 对模型实际使用的短上下文不够困难"],
            ["收益未落在 unseen", "A1 Local-3 seen +1.82 pp，unseen −1.54 pp", "现有提升更像正则化/Stage 2 受益，而非 unseen-order 泛化"],
            ["图合法但经验稀有", "changed views 22.5–24.2% 转移未见，logP/edge 明显下降", "graph 约束不足以保证训练分布内真实性"],
            ["有效覆盖受 gating/fallback 限制", "DualPos 70.2% 样本始终 unchanged", "大部分训练样本从未接受新的顺序监督"],
            ["刷新不是缺少多样性", "可增强样本 Every10 唯一率 84–98%", "Every10 的问题更可能是质量与优化稳定性"],
        ],
        [2200, 3500, 3660],
        font_size=8.2,
        first_col_bold=True,
        caption="表 13  三步排查的联合证据",
    )
    add_heading(doc, "4.2 原因优先级", 2)
    causes = [
        ("第一优先：合法换序对模型来说语义等价或差异过弱", "Graph-valid 只在独立/并行动作之间变化，Atomic Stage 2 主链保持；模型可能只需知道已完成动作集合、最近动作和 Stage。"),
        ("第二优先：Direct Fusion 的顺序表达能力有限", "单 current query 对带位置的历史 token 做集合注意力，没有 history self-attention；true-position control 的完全不变验证了这一点。"),
        ("第三优先：DualPos 抵消了增强干预", "DualPos 既保留真实 recency 又编码 displacement；测试时合法换序几乎不动输出，表明 shift 没有迫使模型学习增强呈现顺序。"),
        ("第四优先：图约束与经验/物理真实性不一致", "Stage 1 多个动作被视为可交换，导致数学合法但训练中少见的邻接；这些视图可能更像 distribution shift 而非有益增强。"),
        ("第五优先：局部测试新颖性与统计功效不足", "Local-3 unseen 仅 134 个 unique clip，Fault unseen 仅 30 个；四位参与者不足以稳定估计小幅泛化收益。"),
        ("第六优先：增强暴露与优化策略", "DualPos 70% 样本 unchanged；Every10 虽然多样，但不断替换视图且不 replay 旧增强，可能产生目标漂移。"),
    ]
    for index, (title, text) in enumerate(causes, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(title + "。")
        set_run_font(r, 11, bold=True, color=DARK_BLUE)
        r = p.add_run(text)
        set_run_font(r, 11)

    add_heading(doc, "4.3 对当前实验结论的修正表述", 2)
    add_callout(
        doc,
        "建议表述",
        "在完整 run 顺序隔离条件下，A1-Legacy-Once 提供了小幅平均收益，但收益未集中在 Local-3 unseen 测试样本。反事实测试显示现有模型对 graph-valid 换序仅弱敏感，DualPos 几乎不敏感。增强历史满足已编码图约束，却包含较高比例的经验未见转移。因此，当前证据更支持‘有限正则化作用’，尚不足以证明方法显著提升未见合法动作顺序的泛化。",
        "info",
    )
def add_recommendations(doc: Document) -> None:
    add_heading(doc, "5. 下一步建议", 1)
    add_para(doc, "前三步排查已经把问题从‘是否有效’缩小为三个可验证方向：模型表达、局部新颖性和增强真实性。建议按以下顺序继续：")
    add_heading(doc, "5.1 先做干净的 2×2 因果消融", 2)
    add_table(
        doc,
        ["因子", "水平 1", "水平 2", "控制要求"],
        [
            ["Shuffle scope", "Broad non-tail prefix", "Active-tail gated / local-only", "同一随机视图集合、相同 changed rate 尽量匹配"],
            ["Position semantics", "Presented position", "True recency + displacement", "同 backbone、features、seeds、epochs"],
        ],
        [1800, 2600, 2800, 2160],
        font_size=8.8,
        first_col_bold=True,
        caption="表 14  建议的 scope × position 2×2 消融",
    )
    add_para(doc, "该实验可以把 A1 与 DualPos 当前混在一起的两个变化拆开：究竟是 active gating 减少了覆盖，还是 true recency 使模型绕过 shuffle。")
    add_heading(doc, "5.2 给 sampler 加经验真实性约束", 2)
    add_bullets(doc, [
        "优先使用训练折 transition likelihood 对候选排序或过滤；限制 novel-transition fraction 与 logP 下降。",
        "对 Stage 1 增加设备状态、cover/pedal/lock 状态和 must-immediately-previous 的显式规则，而不只依赖 feature-history all_must_previous。",
        "生成完整 coherent run 或至少在同一 run 内共享一致的增强前缀，避免每个 current clip 独立采样导致前后不一致。",
        "记录并控制 changed rate、Kendall、最近 token 移动率和真实 prefix 距离，使 A1 / DualPos 强度可比。",
    ])
    add_heading(doc, "5.3 让评估真正针对合法 unseen order", 2)
    add_bullets(doc, [
        "在现有测试中单独报告 Local-3 unseen、transition unseen、exact prefix unseen 和 order-sensitive nodes，而不是只看 All。",
        "构造人工/专家确认的合法替代顺序压力集：同一动作集合、不同合法顺序，并保证每个顺序的设备状态可达。",
        "以 participant 为独立统计单位；seed 只用于估计训练波动，不把 12 个 job 当作 12 个独立参与者。",
    ])
    add_heading(doc, "5.4 改进模型顺序表达的最小实验", 2)
    add_bullets(doc, [
        "增加 history self-attention 或轻量 temporal encoder，使 token 之间可以交互，而非只有 current query 对集合做池化。",
        "加入相对位置/相邻转移编码，并用 valid-vs-corrupted order 辅助任务确认模型确实能区分顺序。",
        "继续保留本报告的 counterfactual test 作为训练后的必测项：若 graph-valid changed-only top-1 仍接近 0，则无需继续扩大 shuffle 次数。",
    ])
    add_callout(
        doc,
        "停止条件",
        "下一轮只有在‘模型对专家确认的合法换序有可测敏感性’且‘增强视图的经验/状态真实性达到预设阈值’后，才值得继续比较 Once 与 Every10。否则增加 refresh 只会放大当前分布偏移。",
        "risk",
    )
def add_appendix(doc: Document) -> None:
    add_heading(doc, "附录 A：可复核文件", 1)
    add_table(
        doc,
        ["文件", "内容"],
        [
            ["order_sensitivity_samples.csv", "所有 checkpoint × sample × counterfactual view 的逐样本输出差异"],
            ["order_sensitivity_summary.csv", "eligible 与 graph-permutation-changed-only 汇总"],
            ["test_group_metadata.csv", "Local-1/2/3、exact prefix、history length、Stage、active tail 标签"],
            ["grouped_performance_summary.csv", "Normal / Fault / All 的分组性能"],
            ["grouped_performance_model_deltas_summary.csv", "A1 / DualPos 对 M2 的配对 fold×seed 差值"],
            ["augmentation_history_audit_detailed.csv", "103,356 个增强视图的结构、经验、扰动明细"],
            ["augmentation_diversity_by_sample.csv", "每个训练样本跨 seed / refresh 的唯一性与重复度"],
            ["manual_audit_candidates.csv", "31 个动作标签人工审查候选"],
            ["manual_semantic_review.md", "人工审查口径、判断和代表性 before/after 案例"],
            ["diagnostic_results.json", "报告使用的紧凑机器可读汇总"],
        ],
        [3550, 5810],
        font_size=8.5,
        first_col_bold=True,
        caption="表 A1  排查结果文件",
    )
    add_heading(doc, "附录 B：关键实现细节", 1)
    add_bullets(doc, [
        "为避免重新训练，直接加载现有 `last.pth` 与 `test_all.pt`；由于当前执行环境没有 PyTorch，使用只读的 torch ZIP 解析器将 tensor storage 映射为 NumPy，并严格复现 DirectHistoryClassifier 的 Linear、LayerNorm、MultiheadAttention、Fusion 与 Node classifier。",
        "原始 actual-order top-1 与项目保存的预测 CSV 在 36/36 checkpoint、全部测试样本上 100% 一致。该验证只保证 top-1 复现；所有 counterfactual 比较均由同一 NumPy 前向路径生成，因此差值不受混合推理后端影响。",
        "Tier3 概率由 35 个 Node 概率按 task graph 的 node-to-tier3 映射求和；与原训练评估逻辑一致。",
        "Graph-valid test permutation 使用原 `augment_history`、`TaskGraph` 和 deterministic stable seed；未查看或定向匹配测试顺序。",
        "分组性能为 12 个 fold×seed job 的非加权 mean ± sample SD；seed 不是独立 participant，报告不据此给出正式显著性结论。",
    ])
    add_heading(doc, "附录 C：局限性", 1)
    add_bullets(doc, [
        "测试顺序敏感性对每个 sample/seed 只抽取一个 graph-valid counterfactual；它估计典型反应，不代表该 history 的最大可能反应。",
        "经验真实性使用一阶相邻转移与真实 prefix 匹配，无法完整描述高阶状态；某些新邻接可能物理可行，只是训练折未观察到。",
        "人工审查基于动作标签，不包含视频与传感器状态，也未由独立工艺专家盲审。",
        "LOSO 仅 4 位参与者，participant-level 推断功效有限；Fault unseen 的 unique n 很小。",
        "分组结果来自分别训练的模型，不能替代同 checkpoint 的因果输入消融；其作用是定位效应集中区域。",
    ])
    add_heading(doc, "附录 D：最终回答", 1)
    add_callout(
        doc,
        "一句话结论",
        "Graph-valid shuffle 的平均作用不显著，主要不是因为完全没有生成新顺序，而是因为测试局部顺序大多已见、Direct Fusion 对合法换序弱敏感、DualPos 保留真实顺序信息，并且部分生成顺序虽满足图约束却缺少经验真实性。",
        "info",
    )


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Graph-Valid Shuffle Augmentation 前三步排查报告"
    props.subject = "Sequence-disjoint experiment diagnostic analysis"
    props.author = "Codex"
    props.keywords = "graph-valid shuffle, atomic-tail, sequence-disjoint, order sensitivity"
    props.comments = "Generated from completed experiment outputs; 2026-08-25"


def main() -> None:
    data = load_data()
    doc = Document()
    configure_styles(doc)
    set_core_properties(doc)
    add_title_page(doc)
    add_executive_summary(doc, data)
    add_method_one(doc, data)
    add_method_two(doc, data)
    add_method_three(doc, data)
    add_integrated_diagnosis(doc)
    add_recommendations(doc)
    add_appendix(doc)
    doc.save(OUT)
    print(str(OUT).encode("unicode_escape").decode("ascii"))


if __name__ == "__main__":
    main()
