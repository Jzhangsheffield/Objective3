from pathlib import Path
from datetime import date
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Thermal_Crimping_TaskGraph_Paper_Draft_2026-08-19.docx"
ASSETS = ROOT / "assets"
ASSETS.mkdir(parents=True, exist_ok=True)

NAVY = "16324F"
BLUE = "276C9E"
TEAL = "2D7C78"
LIGHT_BLUE = "EAF2F8"
PALE = "F4F7F9"
GOLD = "C7932E"
RED = "A53A3A"
GREEN = "2D6A4F"
GREY = "5F6B75"
WHITE = "FFFFFF"
BLACK = "1F2933"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if value and keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    elif not value and keep is not None:
        p_pr.remove(keep)


def set_cell_text(cell, text, bold=False, color=BLACK, size=8.8, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    return p


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    run.font.size = Pt(8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_run_font(run, name="Calibri", size=10.5, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_text(doc, text="", style=None, bold_lead=None, italic=False, align=None, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)
    if keep:
        set_keep_with_next(p)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.65 + 0.45 * level)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=10.2)
    return p


def add_numbered(doc, text, level=0):
    style = "List Number" if level == 0 else "List Number 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.65 + 0.45 * level)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=10.2)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.22)
    table.columns[1].width = Cm(16.2)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), fill)
    table.cell(0, 0).text = ""
    c = table.cell(0, 1)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.1, color=accent, bold=True)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.6)
    set_cell_margins(c, top=130, bottom=130, start=160, end=160)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths=None, font_size=8.5, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color=WHITE, size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], header_fill)
        if widths:
            hdr.cells[i].width = widths[i]
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if ridx % 2 == 0 else PALE
        for i, val in enumerate(row):
            p = set_cell_text(cells[i], val, size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(cells[i], fill)
            if widths:
                cells[i].width = widths[i]
        tr_pr = table.rows[-1]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=8.8, color=GREY, italic=True)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, name="Cambria Math", size=10.5)
    return p


def _font(size, bold=False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    for candidate in [Path("C:/Windows/Fonts") / name, Path("C:/Windows/Fonts/calibri.ttf")]:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _multiline_center(draw, box, text, font, fill="#1F2933", spacing=6):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2), text, font=font, fill=fill, spacing=spacing, align="center")


def _arrow(draw, start, end, fill="#5F6B75", width=5):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    spread = 0.55
    p1 = (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread))
    p2 = (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread))
    draw.polygon([end, p1, p2], fill=fill)


def make_pipeline_figure(path):
    img = Image.new("RGB", (1900, 780), "white")
    draw = ImageDraw.Draw(img)
    draw.text((55, 35), "Training and inference pipeline", font=_font(31, True), fill="#16324F")
    boxes = [
        (45, 175, 290, 315, "Current clip\n16 RGB frames", "#EAF2F8", "#276C9E"),
        (45, 440, 290, 580, "Causal history\nearlier clips only", "#EEF7F4", "#2D7C78"),
        (370, 175, 650, 315, "Frozen R3D-18\n512-D feature", "#F4F7F9", "#5F6B75"),
        (370, 440, 650, 580, "Atomic-tail\ngraph-valid order", "#FFF6E2", "#C7932E"),
        (760, 290, 1040, 465, "Position-aware\nattention", "#EAF2F8", "#276C9E"),
        (1145, 290, 1435, 465, "Direct residual\nfeature fusion", "#EEF7F4", "#2D7C78"),
        (1540, 290, 1850, 465, "35 process nodes\n31 action classes", "#F2EDF8", "#6D4C91"),
    ]
    for x1, y1, x2, y2, label, fc, ec in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fc, outline=ec, width=5)
        _multiline_center(draw, (x1, y1, x2, y2), label, _font(28, True))
    for a, b in [
        ((290, 245), (370, 245)), ((290, 510), (370, 510)),
        ((650, 245), (760, 340)), ((650, 510), (760, 415)),
        ((1040, 378), (1145, 378)), ((1435, 378), (1540, 378)),
    ]:
        _arrow(draw, a, b)
    draw.text((370, 680), "Graph labels select valid training orders only; inference uses actual chronological RGB history.", font=_font(25), fill="#5F6B75")
    img.save(path)


def make_results_figure(path):
    labels = ["M0\nclip-only", "M1\nhistory", "M2\n+position", "M3\nfixed shuffle", "Atomic-tail\nonce"]
    node = [69.81, 79.99, 90.57, 90.05, 91.03]
    action = [83.32, 84.97, 90.64, 90.27, 91.18]
    img = Image.new("RGB", (1800, 830), "white")
    draw = ImageDraw.Draw(img)
    left, top, right, bottom = 150, 135, 1730, 660
    draw.text((150, 30), "Strict LOSO accuracy under actual chronological test history", font=_font(31, True), fill="#16324F")
    draw.rectangle((1110, 47, 1138, 75), fill="#276C9E")
    draw.text((1150, 47), "Process-node accuracy", font=_font(22), fill="#1F2933")
    draw.rectangle((1450, 47, 1478, 75), fill="#2D7C78")
    draw.text((1490, 47), "Action-class accuracy", font=_font(22), fill="#1F2933")
    for value in range(60, 96, 5):
        y = bottom - (value - 60) / 35 * (bottom - top)
        draw.line([(left, y), (right, y)], fill="#D7DEE5", width=2)
        label = f"{value}"
        draw.text((92, y - 14), label, font=_font(22), fill="#5F6B75")
    draw.line([(left, top), (left, bottom)], fill="#5F6B75", width=3)
    draw.line([(left, bottom), (right, bottom)], fill="#5F6B75", width=3)
    group = (right - left) / 5
    bar_w = 74
    for i, (lab, nv, av) in enumerate(zip(labels, node, action)):
        cx = left + group * (i + 0.5)
        for value, color, dx in [(nv, "#276C9E", -bar_w - 5), (av, "#2D7C78", 5)]:
            y = bottom - (value - 60) / 35 * (bottom - top)
            draw.rectangle((cx + dx, y, cx + dx + bar_w, bottom), fill=color)
            txt = f"{value:.2f}"
            bb = draw.textbbox((0, 0), txt, font=_font(21, True))
            draw.text((cx + dx + (bar_w - (bb[2]-bb[0]))/2, y - 30), txt, font=_font(21, True), fill=color)
        _multiline_center(draw, (cx - 115, bottom + 20, cx + 115, bottom + 125), lab, _font(22, True), spacing=4)
    img.save(path)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.15)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Title", 25, NAVY, 0, 10),
        ("Subtitle", 12, GREY, 0, 6),
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 12.5, BLUE, 12, 5),
        ("Heading 3", 10.8, TEAL, 8, 3),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "Small Note" not in [s.name for s in styles]:
        small = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        small.font.name = "Calibri"
        small._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        small.font.size = Pt(8.5)
        small.font.color.rgb = RGBColor.from_string(GREY)
        small.paragraph_format.space_after = Pt(4)
        small.paragraph_format.line_spacing = 1.05

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("投稿可行性评估与论文初稿")
    set_run_font(r, name="Microsoft YaHei", size=27, color=NAVY, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(8)
    r = p2.add_run("基于 TaskGraph Shuffle 与 Atomic-Tail 的热压接细粒度动作识别")
    set_run_font(r, name="Microsoft YaHei", size=14, color=BLUE, bold=True)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(22)
    r = p3.add_run("Journal-readiness assessment · Target-journal shortlist · English manuscript draft")
    set_run_font(r, size=10.5, color=GREY, italic=True)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = line.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(rr, size=11, color=GOLD)

    t = doc.add_table(rows=4, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    values = [
        ("建议主投", "IEEE Access（完成关键补强后）"),
        ("稿件定位", "因果历史融合为主贡献；Atomic-tail 为次级训练增强"),
        ("任务边界", "预切分动作片段识别；明确不含实时动作边界检测"),
        ("版本日期", "19 August 2026"),
    ]
    for i, (a, b) in enumerate(values):
        set_cell_text(t.cell(i, 0), a, bold=True, color=NAVY, size=9.2)
        set_cell_text(t.cell(i, 1), b, size=9.2)
        set_cell_shading(t.cell(i, 0), LIGHT_BLUE)
        set_cell_shading(t.cell(i, 1), WHITE)
        t.cell(i, 0).width = Cm(3.1)
        t.cell(i, 1).width = Cm(11.4)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(35)
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p4.add_run("Prepared from the project’s strict LOSO results and atomic-tail audit artifacts")
    set_run_font(r, size=8.5, color=GREY)
    doc.add_page_break()


def add_assessment(doc):
    doc.add_heading("第一部分｜投稿可行性评估", level=1)
    add_callout(
        doc,
        "结论先行",
        "该工作与 IEEE Access 的应用工程与智能制造范围匹配，但当前证据强度尚不足以直接形成一篇稳健的 IEEE Access 投稿。建议先完成数据集说明、伦理与复现材料、外部或更强基线、层级统计检验四类关键补强，再以 IEEE Access 为首选。现阶段更准确的状态是“Major revision before submission”，而不是“ready to submit”。",
        fill="EEF4F8",
        accent=NAVY,
    )

    doc.add_heading("1.1 审稿视角下的成熟度", level=2)
    readiness_rows = [
        ("期刊范围匹配", "4.5 / 5", "工业动作识别、视频理解、制造过程知识与工程应用高度匹配。"),
        ("技术正确性", "4.0 / 5", "严格跨人 LOSO、三随机种子、无未来信息、participant-first 汇总，实验纪律较好。"),
        ("方法新颖性", "3.0 / 5", "位置感知历史融合价值清晰；atomic-tail 增益小且跨种子不稳定，不能单独支撑强创新叙事。"),
        ("证据广度", "2.5 / 5", "仅 4 位参与者、单一视角、单一内部数据集，缺乏外部基准与域外验证。"),
        ("可复现与数据治理", "2.5 / 5", "代码与审计较完整，但尚缺正式数据说明、伦理/同意、人口统计、标注一致性和公开计划。"),
        ("总体", "3.2 / 5", "具备论文骨架和可信主结果；完成关键补强后可合理冲击 IEEE Access。"),
    ]
    add_table(doc, ["维度", "评分", "判断依据"], readiness_rows, [Cm(3.0), Cm(2.0), Cm(11.2)], font_size=8.8)

    doc.add_heading("1.2 最适合的论文主张", level=2)
    add_text(doc, "建议把论文的中心命题写成：在细粒度工业流程中，单个动作片段往往无法区分具有相似外观但处于不同流程位置的节点；只使用过去片段的因果历史、位置编码与流程图约束，可以显著提升跨操作者的流程节点识别。Atomic-tail 的价值是保留正在执行的原子序列尾部，同时对其余合法历史顺序做受限扰动，从而研究模型对流程顺序变化的鲁棒性。")
    add_bullet(doc, "主贡献：位置感知的因果历史直接特征融合。M2 Direct 相比仅看当前片段的 M0，流程节点准确率提高 20.76 个百分点（69.81% → 90.57%），31 类动作准确率提高 7.32 个百分点（83.32% → 90.64%）。")
    add_bullet(doc, "次贡献：Atomic-tail graph-valid training augmentation。最佳 once 策略达到 91.03% 节点准确率与 91.18% 动作准确率，但相对 M2 Direct 仅 +0.46 / +0.54 个百分点。")
    add_bullet(doc, "可靠性贡献：严格 leave-one-subject-out、每折三随机种子、participant-first 汇总、无验证集调参、仅用过去历史、实际时间顺序测试，以及 72 个原子尾部审计零违规。")
    add_bullet(doc, "应用贡献：热压接流程包含重复视觉动作、正常与故障运行及多模态同步记录，为制造过程节点识别提供了具有真实流程约束的案例。")
    add_callout(doc, "不能写成的主张", "不要声称 atomic-tail 获得了统计显著提升：其 12 个参与者-种子配对结果为 6 胜 6 负，描述性 95% 区间跨越 0。也不要把 fault-run 分类结果表述为“故障检测”，更不要纳入实时边界检测结果。", fill="FBECEC", accent=RED)

    doc.add_heading("1.3 期刊推荐与投稿梯队", level=2)
    journal_rows = [
        ("IEEE Access", "首选；补强后投稿", "范围非常匹配，接受应用工程与制造技术；当前官方页面列出 IF 4.2、平均接收率 20%、4–6 周、APC USD 2,160。", "需补伦理/数据治理、外部或强基线、层级统计与复杂度。"),
        ("Sensors", "并列可行", "若突出三相机/可穿戴传感体系、同步数据集与传感融合潜力，主题合适；官方页面当前列出 IF 4.0，APC CHF 2,600。", "仅使用单路 RGB 会削弱“sensor”叙事；最好加入多视角或 EMG/IMU 消融。"),
        ("Machine Vision and Applications", "更稳妥", "工业机器视觉与工程实现契合；对单一工业案例相对友好。", "仍需与现代视频骨干/程序活动方法比较，并清楚限制泛化。"),
        ("JVCIR", "可行备选", "视频理解、视觉表征与深度学习范围匹配；当前页面列出 IF 3.1。", "更看重视觉方法贡献，建议加入 VideoMAE/Video Swin 等强特征基线或外部数据。"),
        ("Journal of Intelligent Manufacturing", "进取型", "制造业 HAR 与人机协作高度匹配，已有 Praxis 等相近案例。", "门槛高于当前版本；需更多参与者、系统级价值或部署验证。"),
        ("Computers in Industry", "暂不建议首投", "制造 ICT 与工业数据系统匹配，但强调可推广的工业信息技术贡献。", "当前单数据集、4 人、单视角证据不足；需外部验证与更强一般化结论。"),
    ]
    add_table(doc, ["期刊", "建议", "为什么适合", "提交前要求"], journal_rows, [Cm(2.7), Cm(2.2), Cm(6.0), Cm(5.4)], font_size=7.9)

    p = doc.add_paragraph(style="Small Note")
    p.add_run("期刊信息核对（2026-08-19）：")
    add_hyperlink(p, "IEEE Access 官方 About", "https://ieeeaccess.ieee.org/about/")
    p.add_run("；")
    add_hyperlink(p, "IEEE Access APC", "https://ieeeaccess.ieee.org/about/article-processing-charges/")
    p.add_run("；")
    add_hyperlink(p, "Sensors", "https://www.mdpi.com/journal/sensors")
    p.add_run("；")
    add_hyperlink(p, "JVCIR", "https://www.sciencedirect.com/journal/journal-of-visual-communication-and-image-representation")
    p.add_run("；")
    add_hyperlink(p, "Computers in Industry", "https://www.sciencedirect.com/journal/computers-in-industry")
    p.add_run("。影响因子与费用会变化，投稿前须再次核对。")

    doc.add_heading("1.4 现有证据中最有说服力的部分", level=2)
    strength_rows = [
        ("历史上下文确实解决节点歧义", "M2 Direct vs M0：节点 +20.76 pp；动作 +7.32 pp；103/103 条测试运行的节点准确率均改善。"),
        ("位置编码是关键", "M2 Direct vs M1 Direct：节点 +10.57 pp；动作 +5.67 pp；12/12 参与者-种子配对均为正。"),
        ("直接特征融合优于 logits 增量", "M2 Direct vs M2 delta：节点 +6.41 pp；动作 +5.65 pp；节点指标 12/12 为正。"),
        ("流程图增强经过因果审计", "72 个训练单元；atomic-tail 适用比例约 69.39%（all-runs）/ 73.41%（normal-only）；尾部违规为 0。"),
        ("测试顺序不是增益来源", "同一 atomic-tail once 检查点在实际顺序与 atomic 顺序测试的差异 <0.22 pp；仅 39/5,685 个节点预测改变。"),
        ("错误分析可形成讨论", "496/5,685 个错误；仅 8 个是同一 Tier-3 内部的节点混淆；46.8% 错误置信度 ≥0.9，提示剩余瓶颈主要是视觉语义而非流程位置。"),
    ]
    add_table(doc, ["证据点", "现有结果"], strength_rows, [Cm(5.0), Cm(11.2)], font_size=8.6)

    doc.add_heading("1.5 关键不足、风险与改进方法", level=2)
    gaps = [
        ("P0", "参与者数量过少", "4 位参与者不足以支持广泛跨操作者泛化；单个参与者 M 上 atomic-tail 还出现负增益。", "优先增加到至少 10–15 位；若无法追加，采用按参与者层级 bootstrap，并把结论限定为 proof-of-concept。"),
        ("P0", "缺伦理与同意信息", "项目文件中未找到伦理审批、知情同意、隐私处理或参与者人口统计。", "补充伦理编号/豁免、书面同意、面部与生理信号治理、年龄/性别/经验范围；否则数据集稿件风险很高。"),
        ("P0", "缺外部/公开基准", "只在内部热压接数据上验证，难以证明方法可迁移。", "至少在 HA-ViD、Assembly101、HA4M 或 IMPACT 的可比预切分识别任务上复现历史融合；若标签结构不兼容，构建可复现的程序节点子任务。"),
        ("P0", "Atomic-tail 统计证据弱", "+0.46 / +0.54 pp，6 胜 6 负；不同随机种子方向不一致。", "将其定位为受限正则化与鲁棒性分析；扩大种子或参与者，报告层级置信区间与效应量，不使用“significant improvement”。"),
        ("P1", "基线不够现代", "R3D-18 从头训练及有限融合基线，审稿人会要求现代预训练视频模型。", "增加 VideoMAE v2、Video Swin、SlowFast 或 TimeSformer 的冻结/微调基线，并统一 LOSO、输入帧数与计算预算。"),
        ("P1", "多模态数据但只用单路 RGB", "三相机和双侧 EMG/IMU 是数据亮点，但论文实验没有利用。", "至少加入多视角 late fusion；若资源允许，再加 RGB+EMG/IMU。否则在标题与贡献中明确限定 single-view RGB。"),
        ("P1", "缺标注质量指标", "未找到双人标注、复核流程、边界/类别一致性或 κ/IoU。", "抽取 10–20% 双重标注，报告 Cohen’s κ 或 Krippendorff’s α、边界偏差及仲裁流程。"),
        ("P1", "无验证集与超参选择说明", "使用最后一轮避免测试调参是优点，但无法解释超参数如何确定。", "采用训练参与者内部的 nested LOSO/leave-one-run-out 验证，或声明超参数在独立先导集上预定义并冻结。"),
        ("P1", "缺效率与复杂度", "没有参数量之外的训练/推理时间、FLOPs、显存和历史长度敏感性。", "报告 949,027 个融合层参数、总参数、FLOPs、每片段延迟、显存；做历史长度 1/5/10/全部消融。"),
        ("P2", "训练增强依赖历史真值节点", "atomic-tail 用训练样本的历史节点标注选择受保护尾部；若不说明，容易被误解为推理泄漏。", "明确这是仅训练时的监督式增强；推理只输入既往 RGB 特征且采用实际时间顺序，不输入历史真值节点。"),
        ("P2", "故障运行定义可能被误读", "fault-run 结果只是动作识别在含偏差执行中的子集表现。", "提供故障类型与占比；坚持称为 recognition under fault-containing runs，不宣称 fault detection。"),
    ]
    add_table(doc, ["优先级", "不足", "审稿风险", "可执行改进"], gaps, [Cm(1.2), Cm(3.0), Cm(5.2), Cm(7.2)], font_size=7.6)

    doc.add_heading("1.6 建议的最小投稿包", level=2)
    for item in [
        "补齐伦理/同意、参与者信息、采集硬件、同步方法、标注协议与数据可用性声明。",
        "新增至少两个现代视频骨干，以及历史长度、图约束组件和训练顺序策略的统一消融。",
        "采用 participant-level bootstrap 或层级 bootstrap，报告 95% CI、配对效应量和多重比较策略。",
        "加入复杂度与推理效率表；公开严格 LOSO 清单、TaskGraph JSON、随机种子、特征缓存生成方式和训练脚本。",
        "若不扩充参与者，至少在一个公开工业/程序活动数据集上验证核心历史融合模块。",
        "将实时边界检测完全从本文任务、方法、结果和贡献中剥离；仅在局限性中说明本文假设动作片段边界已给定。",
    ]:
        add_numbered(doc, item)


def add_manuscript(doc, pipeline_path, results_path):
    doc.add_page_break()
    doc.add_heading("第二部分｜English Manuscript Draft", level=1)
    add_callout(doc, "Draft status", "This is a content-complete first draft built from the current project artifacts. Bracketed items must be completed before submission. The manuscript deliberately excludes all real-time action-boundary detection experiments and assumes pre-segmented clips.", fill="FFF6E2", accent=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Causal Task-Graph History Fusion with Atomic-Tail Augmentation for Fine-Grained Action Recognition in Thermal Crimping")
    set_run_font(r, size=20, color=NAVY, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("[Author names] · [Affiliations] · [Corresponding author email]")
    set_run_font(r, size=9.5, color=GREY)

    doc.add_heading("Abstract", level=2)
    abstract = (
        "Fine-grained action recognition in industrial workflows is difficult when visually similar manipulations occur at multiple process positions. We study this problem in a thermal-crimping workflow containing 103 executions from four participants, including 76 normal and 27 fault-containing runs. The curated benchmark comprises 1,895 pre-segmented RGB action clips annotated as 35 process nodes and mapped to 31 action categories. We propose a causal history-fusion model that combines a frozen 3D ResNet representation of the current clip with position-aware attention over earlier clips from the same run. We further introduce atomic-tail graph-valid augmentation, a training-only strategy that preserves the valid unfinished prefix of an active atomic sequence while shuffling the remaining history under task-graph constraints. Evaluation uses strict four-fold leave-one-subject-out testing, three random seeds, no test-subject validation, and participant-first aggregation. The position-aware history model improves process-node accuracy from 69.81±3.94% to 90.57±2.66% and action-category accuracy from 83.32±5.66% to 90.64±2.64%. Atomic-tail augmentation with one fixed training draw reaches 91.03±3.13% node accuracy and 91.18±3.06% action accuracy under actual chronological test history. The latter gains over the unaugmented history model are modest (+0.46 and +0.54 percentage points) and vary across participants and seeds, so we interpret the augmentation as a constrained regularizer rather than a uniformly superior estimator. These results show that causal workflow history is the dominant source of improvement for distinguishing repeated industrial actions, while aggressive order perturbation can reduce performance."
    )
    add_text(doc, abstract)
    add_text(doc, "Index Terms—industrial action recognition, procedural video understanding, task graph, causal history, workflow position, data augmentation, thermal crimping.", italic=True)

    doc.add_heading("1. Introduction", level=2)
    add_text(doc, "Human-centered manufacturing increasingly depends on visual systems that can understand not only what motion is visible, but also where that motion occurs within a procedure. This distinction matters in workflows such as thermal crimping, where the same primitive manipulation may be repeated on different components or at different stages. A clip-only recognizer can therefore assign the correct semantic action while confusing the corresponding process node. Such errors limit the usefulness of action recognition for work-in-progress monitoring, instruction support, traceability, and downstream quality assurance.")
    add_text(doc, "Industrial action datasets and procedural benchmarks have advanced from isolated gestures to multi-view assembly, error-containing executions, and structured process annotations [1]–[5]. Yet two practical questions remain underexplored. First, how much does strictly causal visual history help when current-clip appearance is ambiguous? Second, can a known task graph support order augmentation without destroying locally meaningful action sequences? These questions differ from temporal action segmentation: the present work assumes that an action clip has already been segmented and predicts its process-node and action-category labels. We do not evaluate online boundary detection.")
    add_text(doc, "We address these questions with a two-stage architecture. A 3D ResNet-18 encodes each RGB clip. A lightweight direct-fusion head then attends from the current feature to features of earlier clips, using an explicit history-position embedding. This design is causal by construction: history is restricted to earlier annotation rows from the same run and participant. To study robustness to valid procedural variation, we derive graph-valid history permutations from a process task graph. Our atomic-tail variant detects an unfinished atomic sequence using observed history only, keeps its prefix contiguous at the end of the history, and shuffles the remaining clips subject to graph constraints.")
    add_text(doc, "The main contributions are:")
    add_bullet(doc, "A causal, position-aware history-fusion architecture for fine-grained process-node recognition that preserves the current-clip representation through an identity-safe residual initialization.")
    add_bullet(doc, "Atomic-tail graph-valid augmentation, which protects the locally ordered prefix of an active atomic sequence while perturbing the remainder of training history without accessing the current target or future clips.")
    add_bullet(doc, "A strict cross-subject evaluation on 1,895 clips from 103 thermal-crimping runs, with normal and fault-containing executions, three random seeds, participant-first reporting, causal audits, and detailed stage and confusion analyses.")
    add_bullet(doc, "Empirical evidence that causal history and its positional order, rather than unrestricted graph-valid shuffling, drive the largest gains for repeated industrial actions.")

    doc.add_heading("2. Related Work", level=2)
    doc.add_heading("2.1 Video action recognition", level=3)
    add_text(doc, "Modern video recognition models range from 3D convolutional networks such as I3D and SlowFast to transformer-based architectures such as TimeSformer and Video Swin Transformer [6]–[9]. These models are effective at learning clip-level appearance and motion. However, clip-level backbones do not explicitly represent a workflow position, and fine-grained industrial classes can remain ambiguous when the same hand-object interaction occurs repeatedly. Our study therefore treats the video backbone as a controlled feature extractor and isolates the incremental value of causal process history.")
    doc.add_heading("2.2 Industrial and procedural activity datasets", level=3)
    add_text(doc, "Assembly101 provides large-scale multi-view procedural videos with mistakes and corrections [1]. HA4M offers multimodal manufacturing action data [2], while HA-ViD introduces multi-view industrial assembly videos and rich assembly knowledge annotations [3]. Praxis demonstrates an AI-driven action-recognition framework in a manufacturing case study [4]. More recently, IMPACT introduced five-view RGB-D recordings, a prerequisite graph, anomaly–recovery labels, and multi-granularity supervision for real tool-based assembly [5]. Relative to these benchmarks, our dataset is substantially smaller and is not presented as a scale contribution. Its distinctive value is the combination of thermal-crimping process nodes, repeated semantic actions at different workflow positions, normal and fault-containing executions, and synchronized sensing. A public release and complete governance statement are required before claiming a dataset contribution.")
    doc.add_heading("2.3 Temporal and graph-based procedural reasoning", level=3)
    add_text(doc, "Temporal segmentation methods such as MS-TCN model long action sequences [10], while graph-based temporal reasoning explicitly propagates relationships among video units [11]. Task graphs have also been mined from instructional videos to support keystep recognition [12]. Our formulation differs in two ways. First, the prediction target is a pre-segmented clip rather than a framewise boundary-aware sequence. Second, the known engineering task graph is used to constrain training-history permutations, while inference uses the observed chronological history. This isolates recognition from boundary detection and avoids future-context leakage.")

    doc.add_heading("3. Dataset and Task Definition", level=2)
    doc.add_heading("3.1 Thermal-crimping recordings", level=3)
    add_text(doc, "The dataset contains 103 recorded executions performed by four participants (A, D, J, and M). Of these runs, 76 follow the normal workflow and 27 contain one or more recorded deviations. The annotations yield 1,895 pre-segmented action clips: 1,441 from normal runs and 454 from fault-containing runs. Each clip is linked to three synchronized RGB camera paths and wearable recordings from left and right MindRove devices. Each wearable stream contains eight electromyography channels and three-axis accelerometer and gyroscope signals. The experiments in this paper use only camera 001484412812; the remaining views and wearable signals are reserved for future work.")
    add_callout(doc, "Mandatory completion before submission", "Insert camera model, lens/viewpoint, frame rate verified from acquisition metadata, lighting and workstation description, sensor synchronization procedure, participant demographics and experience, recruitment and compensation, ethics approval or exemption number, written consent, privacy handling, and a data/code availability statement.", fill="FBECEC", accent=RED)
    dataset_rows = [
        ("A", "24", "15", "9", "431", "294", "137"),
        ("D", "25", "21", "4", "462", "400", "62"),
        ("J", "30", "21", "9", "555", "387", "168"),
        ("M", "24", "19", "5", "447", "360", "87"),
        ("Total", "103", "76", "27", "1,895", "1,441", "454"),
    ]
    add_table(doc, ["Participant", "Runs", "Normal runs", "Fault runs", "Clips", "Normal clips", "Fault clips"], dataset_rows, [Cm(2.1), Cm(1.5), Cm(2.0), Cm(1.7), Cm(1.8), Cm(2.0), Cm(2.0)], font_size=8.3)
    caption(doc, "Table 1. Dataset distribution used in strict leave-one-subject-out evaluation.")

    doc.add_heading("3.2 Label hierarchy and process graph", level=3)
    add_text(doc, "The workflow contains 35 classified process nodes arranged over three stages. These nodes map to 31 Tier-3 action categories, so several semantically identical actions occur at different process positions. Stage 1 contains 242 clips, Stage 2 contains 1,397 clips, and Stage 3 contains 256 clips. The integrated task graph specifies mandatory predecessors, optional/possible predecessors, immediate predecessors, mutually exclusive relations, and five atomic sequences with lengths 2, 2, 14, 2, and 2. The long Stage-2 atomic sequence captures the central thermal-crimping procedure.")
    add_text(doc, "The two prediction tasks are: (i) process-node recognition over 35 nodes and (ii) action-category recognition over 31 Tier-3 classes. The latter is obtained from the node-to-action mapping. Fault-containing runs are not treated as a fault-detection task; they are evaluated only as a subset on which the same action-recognition labels are predicted.")

    doc.add_heading("3.3 Clip construction", level=3)
    add_text(doc, "For every annotated clip, 16 RGB frames are uniformly sampled, resized to 224×224 pixels, and passed to the video encoder. Across the evaluated view, the 1,895 clips contain 94,724 source frames, with a mean duration of 49.99 frames, median 41 frames, minimum 10 frames, and maximum 653 frames. The original RGB resolution is 1280×720. This paper assumes that the temporal boundaries of each clip are given.")

    doc.add_heading("4. Method", level=2)
    doc.add_picture(str(pipeline_path), width=Inches(6.45))
    caption(doc, "Figure 1. Proposed causal history-fusion pipeline. Task-graph labels are used only to create training orders; evaluation uses actual chronological history.")

    doc.add_heading("4.1 Clip-level video representation", level=3)
    add_text(doc, "A ResNet3D-18 backbone is trained from scratch on the 31 action categories using 16-frame RGB clips. The backbone produces a 512-dimensional representation x for each clip. Backbone training uses cross-entropy loss for 100 epochs, AdamW with an initial learning rate of 10⁻⁴ and weight decay 10⁻⁴, batch size 16, and learning-rate milestones at epochs 50 and 75. After training, clip features are cached and the backbone is frozen for all history-fusion experiments.")

    doc.add_heading("4.2 Causal history construction", level=3)
    add_text(doc, "For a current clip i, its history Hᵢ contains only clips from the same participant and run whose annotation rows precede i. Neither the current target nor any future clip is included. Histories are padded to a maximum length of 35 with a learned null token. This construction is audited before training and is unchanged across model variants.")

    doc.add_heading("4.3 Position-aware direct history fusion", level=3)
    add_text(doc, "The current feature and history features are projected to 256 dimensions. A four-head multi-head attention module uses the current clip as the query and the null-augmented history as keys and values. An optional learned position embedding encodes the ordered history position. The attended context is concatenated with the original current feature and passed through a residual direct-fusion projection:")
    add_equation(doc, "qᵢ = LN(Wq xᵢ),     hᵢ,j = LN(Wh xᵢ,j) + pⱼ")
    add_equation(doc, "cᵢ = MHA(qᵢ, [h∅; hᵢ,1…hᵢ,L], [h∅; hᵢ,1…hᵢ,L])")
    add_equation(doc, "x̃ᵢ = Wf [xᵢ; cᵢ],     ŷᵢ = Wc LN(x̃ᵢ)")
    add_text(doc, "The fusion layer is initialized so that the current-feature branch is an identity mapping and the history branch initially contributes zero. This identity-safe initialization avoids degrading the pretrained clip representation at the start of fusion training. The feature-level head contains 949,027 trainable parameters.")

    doc.add_heading("4.4 Graph-valid and atomic-tail augmentation", level=3)
    add_text(doc, "A graph-valid shuffle creates a topological reordering of the observed history while respecting the process graph. Unrestricted graph-valid perturbation may nevertheless break a locally meaningful sequence near the current step. Atomic-tail augmentation therefore identifies an active atomic sequence using the observed history only. The latest observed node must belong to an atomic sequence, and the observed members of that sequence must form a valid unfinished prefix. This prefix is kept contiguous at the end of the history, while the remaining history is graph-validly shuffled. If a repeated node makes the tail ambiguous, the method falls back to the actual order.")
    add_text(doc, "We compare three refresh policies over 50 fusion epochs: one fixed draw (once), a new draw every 10 epochs, and a new draw every epoch. The once policy exposes each training sample to one augmented history; the other policies expose it to up to five and 50 draws, respectively. Across 72 atomic-tail training units, the audit finds zero protected-tail violations. The method applies to approximately 69.39% of all-run training samples and 73.41% of normal-only training samples.")
    add_callout(doc, "No inference-time label leakage", "Ground-truth historical node IDs select a valid permutation only during supervised training. At inference, the model receives previous RGB features in their actual chronological order and no historical ground-truth node IDs.", fill="EEF7F4", accent=GREEN)

    doc.add_heading("4.5 Optimization", level=3)
    add_text(doc, "All fusion heads are trained for 50 epochs with batch size 64 using AdamW (learning rate 10⁻³, weight decay 10⁻⁴), cross-entropy over 35 process nodes, and gradient clipping at 1.0. The last-epoch checkpoint is evaluated. No validation subject, early stopping, or test-fold tuning is used. Before publication, the rationale or an independent training-only validation protocol for these hyperparameters should be documented.")

    doc.add_heading("5. Experimental Protocol", level=2)
    doc.add_heading("5.1 Strict cross-subject evaluation", level=3)
    add_text(doc, "We use four leave-one-subject-out folds. In each fold, all runs and clips from one participant are held out for testing. The remaining three participants are used for training. Every configuration is repeated with seeds 1, 2, and 42. Metrics are first averaged across the three seeds within each held-out participant and then averaged equally across the four participant means. The reported ± value is the sample standard deviation across the four participant means; the 12 participant-seed pairs are not treated as independent subjects.")
    doc.add_page_break()
    fold_spacer = doc.add_paragraph()
    fold_spacer.paragraph_format.space_after = Pt(24)
    doc.add_heading("5.1.1 Fold composition", level=3)
    loso_rows = [
        ("A", "D, J, M", "431", "35/35", "31/31"),
        ("D", "A, J, M", "462", "35/35", "31/31"),
        ("J", "A, D, M", "555", "35/35", "31/31"),
        ("M", "A, D, J", "447", "35/35", "31/31"),
    ]
    add_table(doc, ["Held-out", "Training participants", "Test clips", "Node coverage", "Action coverage"], loso_rows, [Cm(2.0), Cm(4.5), Cm(2.1), Cm(2.4), Cm(2.4)], font_size=8.5)
    caption(doc, "Table 2. Strict LOSO split. Coverage values refer to the all-runs test set.")

    doc.add_heading("5.2 Compared variants", level=3)
    variant_rows = [
        ("M0", "Current RGB clip only; frozen 512-D feature with node classifier."),
        ("M1 Direct", "Actual causal history; direct feature fusion; no position embedding."),
        ("M2 Direct", "Actual causal history; direct feature fusion; learned history-position embedding."),
        ("M3 Direct", "One fixed graph-valid shuffled history; direct fusion with position embedding."),
        ("Dynamic Direct", "Graph-valid shuffled history refreshed during training."),
        ("Atomic-tail", "Protected atomic prefix plus graph-valid remainder; refresh once, every 10 epochs, or every epoch."),
    ]
    add_table(doc, ["Variant", "Definition"], variant_rows, [Cm(3.1), Cm(12.9)], font_size=8.6)

    doc.add_heading("5.3 Metrics and analysis", level=3)
    add_text(doc, "We report accuracy, macro-F1, and balanced accuracy for both 35 process nodes and 31 action categories. Primary comparisons use all runs and actual chronological test history. We additionally report normal-run and fault-containing-run subsets, stage-wise accuracy, repeated-node confusion groups, immediate versus non-immediate task-graph targets, paired participant-seed differences, and prediction sensitivity to test-history order.")

    doc.add_heading("6. Results", level=2)
    doc.add_picture(str(results_path), width=Inches(6.25))
    caption(doc, "Figure 2. Main all-runs results under strict LOSO and actual chronological test history.")
    result_rows = [
        ("M0 clip-only", "69.81±3.94", "72.91±3.51", "83.32±5.66", "81.41±3.98"),
        ("M1 Direct", "79.99±6.52", "80.69±4.10", "84.97±4.96", "84.04±2.24"),
        ("M2 Direct", "90.57±2.66", "87.81±2.79", "90.64±2.64", "87.06±3.00"),
        ("M3 Direct", "90.05±3.31", "87.60±3.32", "90.27±3.10", "87.23±3.05"),
        ("Dynamic Direct", "89.79", "—", "90.02", "—"),
        ("Atomic every epoch", "89.18", "—", "89.45", "—"),
        ("Atomic every 10 epochs", "89.73", "—", "90.29", "—"),
        ("Atomic once", "91.03±3.13", "88.42±2.70", "91.18±3.06", "87.81±2.58"),
    ]
    add_table(doc, ["Model", "Node Acc.", "Node Macro-F1", "Action Acc.", "Action Macro-F1"], result_rows, [Cm(3.6), Cm(2.8), Cm(3.1), Cm(2.8), Cm(3.1)], font_size=8.1)
    caption(doc, "Table 3. All-runs test performance (%). Values with ± use participant-first mean ± sample SD across four held-out participants. A dash indicates that the current summary artifact did not expose the metric.")

    doc.add_heading("6.1 Causal history is the dominant improvement", level=3)
    add_text(doc, "M2 Direct improves process-node accuracy over M0 by 20.76 percentage points and action-category accuracy by 7.32 points. The larger node gain supports the central hypothesis: history primarily resolves workflow-position ambiguity among actions that are visually or semantically repeated. M2 Direct also improves node accuracy on all 103 test runs relative to M0. Stage 2, which contains the long atomic sequence and most repeated actions, gains 26.72 points in node accuracy. Four major repeated-node confusion groups drop from 166, 247, 206, and 150 M0 errors to 0, 0, 3, and 1 M2 Direct errors, respectively.")

    doc.add_heading("6.2 Position order matters", level=3)
    add_text(doc, "Removing position embeddings (M1 Direct) reduces node accuracy from 90.57% to 79.99% and action accuracy from 90.64% to 84.97%. All 12 participant-seed comparisons favor M2 Direct for both accuracies. Conversely, fixed graph-valid shuffling (M3 Direct) is 0.52 points lower in node accuracy and 0.37 points lower in action accuracy than actual-order M2 Direct. The dynamic shuffle is also lower. These results indicate that graph validity alone does not preserve all discriminative temporal information; the observed chronological position remains useful.")

    doc.add_heading("6.3 Atomic-tail augmentation", level=3)
    add_text(doc, "Atomic-tail once gives the highest all-runs mean: 91.03±3.13% process-node accuracy, 88.42±2.70% node macro-F1, 91.18±3.06% action accuracy, and 87.81±2.58% action macro-F1. Its balanced accuracies are 88.95±2.82% and 88.40±2.77% for nodes and actions. Relative to M2 Direct, however, the gains are only 0.46 and 0.54 points. Across 12 participant-seed pairs, both metrics show six wins and six losses; the descriptive 95% interval for the node difference is [−1.45, +2.37] points. The seed-level node differences are −2.13, +0.43, and +3.09 points for seeds 1, 2, and 42. Therefore, the mean improvement should not be described as statistically significant or uniform.")
    add_text(doc, "Participant-level node differences are +1.16 (A), +0.51 (D), +1.08 (J), and −0.89 (M) points. More frequent refresh is harmful in all-runs evaluation: every 10 epochs reaches 89.73% and every epoch 89.18%, compared with 91.03% for once. This pattern suggests an augmentation-strength trade-off: a single structurally plausible perturbation may regularize the history head, whereas repeated reordering can obscure the empirical temporal distribution.")

    doc.add_heading("6.4 Actual-order evaluation and subset results", level=3)
    add_text(doc, "The final evaluation uses actual chronological test history, identical to M2 Direct. Evaluating the same atomic-tail once checkpoint with atomic-order history changes average accuracy by less than 0.22 points; only 39 of 5,685 node predictions change. Thus, the reported average gain is not an artifact of test-time reordering.")
    subset_rows = [
        ("All runs", "91.03±3.13", "88.42±2.70", "91.18±3.06", "87.81±2.58"),
        ("Normal runs", "91.69", "88.96", "91.80", "88.38"),
        ("Fault-containing runs", "90.19", "86.31", "90.37", "85.20"),
    ]
    add_table(doc, ["Subset", "Node Acc.", "Node Macro-F1", "Action Acc.", "Action Macro-F1"], subset_rows, [Cm(4.1), Cm(2.8), Cm(3.1), Cm(2.8), Cm(3.1)], font_size=8.3)
    caption(doc, "Table 4. Atomic-tail once under actual chronological test history (%). Fault-containing runs are an action-recognition subset, not a fault-detection experiment.")

    doc.add_heading("6.5 Stage and error analysis", level=3)
    add_text(doc, "Atomic-tail once achieves 84.17%, 92.81%, and 87.67% node accuracy in Stages 1, 2, and 3, respectively. Accuracy is 92.58% for immediate task-graph targets and 86.47% for non-immediate targets. The repeated-node group reaches 93.75%, with seven remaining pairwise repeated-node errors across all participant-seed predictions.")
    add_text(doc, "Across 5,685 predictions, 496 are incorrect. Only eight errors (1.6%) confuse process nodes mapped to the same action category, indicating that most remaining errors cross semantic action categories rather than only workflow positions. Moreover, 232 errors (46.8%) have confidence at least 0.9. Frequent confusions include node 24→12 (30), 19→20 (21), 34→24 (19), 24→25 (16), and 18→23 (13). Nodes 34 and 1 have the lowest recalls (61.3% and 65.2%). These findings motivate stronger visual representations, object-aware features, and multi-view fusion.")

    doc.add_heading("7. Discussion", level=2)
    doc.add_heading("7.1 Why causal history works", level=3)
    add_text(doc, "The performance gap between node and action recognition for M0 indicates that a clip often contains enough evidence for a semantic action but not for its exact process location. Causal history supplies the missing state proxy: previous manipulations narrow the set of feasible current nodes. The 10.57-point improvement from adding position embeddings further shows that a bag of previous features is insufficient; their ordered positions encode useful progression information.")

    doc.add_heading("7.2 Why stronger shuffling can hurt", level=3)
    add_text(doc, "A task graph describes valid precedence, but it does not encode every empirical dependency, operator habit, duration pattern, or local visual transition. Repeated graph-valid reordering may therefore create histories that are logically legal yet statistically rare. Atomic-tail mitigates this problem by protecting a locally active sequence, but the negative effect of frequent refresh suggests that the remaining perturbation can still be too strong. Future work should estimate permutation plausibility from training data or weight shuffled histories by their distance from observed orders.")

    doc.add_heading("7.3 Practical interpretation", level=3)
    add_text(doc, "The current model is suitable as an offline recognition component for pre-segmented action clips. It does not establish real-time latency, online boundary detection, fault diagnosis, or deployment safety. The results on fault-containing runs show that recognition remains relatively robust when deviations are present, but they do not identify the deviations themselves. Any manufacturing deployment would require boundary handling, uncertainty calibration, privacy controls, and prospective validation.")

    doc.add_heading("8. Limitations", level=2)
    add_text(doc, "This study has several important limitations. First, only four participants are available, so between-operator uncertainty is estimated from four held-out units. Second, the experiments use one camera although the recordings include three RGB views and wearable signals. Third, all results are obtained on one thermal-crimping workflow; external generalization is untested. Fourth, atomic-tail augmentation uses historical ground-truth process-node labels during training to choose a protected prefix, although inference does not require these labels. Fifth, the current annotation-quality evidence does not include an inter-annotator agreement study. Sixth, hyperparameters were not selected with a nested validation protocol. Seventh, the task assumes known clip boundaries and excludes online temporal boundary detection. Finally, the modest and seed-dependent atomic-tail gain requires confirmation with more participants, stronger backbones, and hierarchical uncertainty estimates.")

    doc.add_heading("9. Conclusion", level=2)
    add_text(doc, "We presented a causal task-graph history-fusion approach for fine-grained action recognition in thermal crimping. Strict cross-subject experiments show that position-aware past context substantially improves process-node recognition over a clip-only model, especially for repeated actions in the central workflow stage. Atomic-tail graph-valid augmentation achieves the best mean accuracy when applied once, but its small and variable gain calls for a conservative interpretation. The clearest conclusion is that chronological causal history is highly informative for industrial workflow-position recognition, while graph-constrained order augmentation must preserve local procedural structure and be carefully calibrated. Broader participant coverage, public benchmarking, modern video backbones, multimodal fusion, and complete data-governance documentation are required before strong claims of general industrial applicability.")

    doc.add_heading("Declarations", level=2)
    add_text(doc, "Ethics approval and consent to participate—[TO COMPLETE: approval/exemption body, reference number, written informed consent, privacy treatment].")
    add_text(doc, "Data availability—[TO COMPLETE: repository, license, controlled-access conditions, or justified non-release statement].")
    add_text(doc, "Code availability—[TO COMPLETE: repository and archived release containing LOSO manifests, task graph, configuration files, seeds, and audit scripts].")
    add_text(doc, "Competing interests—The authors declare [TO COMPLETE].")
    add_text(doc, "Funding—[TO COMPLETE].")
    add_text(doc, "Author contributions—[TO COMPLETE using CRediT taxonomy].")

    doc.add_heading("References", level=2)
    references = [
        "[1] F. Sener et al., ‘Assembly101: A Large-Scale Multi-View Video Dataset for Understanding Procedural Activities,’ Proc. IEEE/CVF CVPR, pp. 21096–21106, 2022.",
        "[2] F. Cicirelli et al., ‘The HA4M dataset: Multi-modal monitoring of an assembly task for human action recognition in manufacturing,’ Scientific Data, vol. 9, art. 745, 2022, doi: 10.1038/s41597-022-01843-z.",
        "[3] H. Zheng, R. Lee, and Y. Lu, ‘HA-ViD: A Human Assembly Video Dataset for Comprehensive Assembly Knowledge Understanding,’ Advances in Neural Information Processing Systems, vol. 36, 2023, doi: 10.52202/075280-2930.",
        "[4] C. Gkournelos et al., ‘Praxis: a framework for AI-driven human action recognition in assembly,’ Journal of Intelligent Manufacturing, vol. 35, pp. 3697–3711, 2024, doi: 10.1007/s10845-023-02228-8.",
        "[5] D. Wen et al., ‘IMPACT: A Dataset for Multi-Granularity Human Procedural Action Understanding in Industrial Assembly,’ arXiv:2604.10409, 2026.",
        "[6] J. Carreira and A. Zisserman, ‘Quo Vadis, Action Recognition? A New Model and the Kinetics Dataset,’ Proc. IEEE CVPR, 2017.",
        "[7] C. Feichtenhofer, H. Fan, J. Malik, and K. He, ‘SlowFast Networks for Video Recognition,’ Proc. IEEE/CVF ICCV, 2019.",
        "[8] G. Bertasius, H. Wang, and L. Torresani, ‘Is Space-Time Attention All You Need for Video Understanding?’ Proc. ICML, 2021.",
        "[9] Z. Liu et al., ‘Video Swin Transformer,’ Proc. IEEE/CVF CVPR, 2022.",
        "[10] Y. A. Farha and J. Gall, ‘MS-TCN: Multi-Stage Temporal Convolutional Network for Action Segmentation,’ Proc. IEEE/CVF CVPR, 2019.",
        "[11] Y. Huang, Y. Sugano, and Y. Sato, ‘Improving Action Segmentation via Graph-Based Temporal Reasoning,’ Proc. IEEE/CVF CVPR, 2020.",
        "[12] K. Zhou et al., ‘Video-Mined Task Graphs for Keystep Recognition in Instructional Videos,’ Advances in Neural Information Processing Systems, vol. 36, 2023.",
    ]
    for ref in references:
        p = doc.add_paragraph(style="Small Note")
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.add_run(ref)
    p = doc.add_paragraph(style="Small Note")
    p.add_run("Reference note: verify all author lists, page ranges, DOIs, and BibTeX against the target journal’s style before submission. Key primary sources: ")
    add_hyperlink(p, "Assembly101", "https://openaccess.thecvf.com/content/CVPR2022/html/Sener_Assembly101_A_Large-Scale_Multi-View_Video_Dataset_for_Understanding_Procedural_Activities_CVPR_2022_paper.html")
    p.add_run("; ")
    add_hyperlink(p, "HA-ViD", "https://proceedings.neurips.cc/paper_files/paper/2023/hash/d40e6e4b3ee6c24f2bf2cb72c2412f4b-Abstract-Datasets_and_Benchmarks.html")
    p.add_run("; ")
    add_hyperlink(p, "IMPACT", "https://arxiv.org/abs/2604.10409")
    p.add_run("; ")
    add_hyperlink(p, "Praxis", "https://link.springer.com/article/10.1007/s10845-023-02228-8")
    p.add_run(".")

    doc.add_page_break()
    doc.add_heading("Appendix A｜Claim and Evidence Audit", level=2)
    audit_rows = [
        ("Causal history strongly improves workflow-node recognition.", "Supported", "M2 Direct vs M0: +20.76 pp node accuracy; positive on 103/103 test runs."),
        ("Position encoding is important.", "Supported", "M2 Direct vs M1 Direct: +10.57 pp node accuracy; 12/12 paired positives."),
        ("Atomic-tail significantly outperforms actual history.", "Not supported", "+0.46 pp, 6 wins/6 losses, interval crosses zero."),
        ("Graph-valid shuffling improves robustness.", "Mixed/negative", "M3 and dynamic variants are below M2 Direct on all-runs means."),
        ("The method detects faults.", "Not evaluated", "Fault-containing runs are recognition subsets only."),
        ("The method works online with unknown boundaries.", "Not evaluated", "All experiments use pre-segmented clips; boundary detection is excluded."),
        ("The method generalizes across operators.", "Preliminary", "Strict LOSO is sound, but only four participants are available."),
        ("The augmentation has no target leakage.", "Supported by audit", "Selection uses observed prior history only; test input uses actual chronological RGB history."),
    ]
    add_table(doc, ["Potential claim", "Status", "Evidence / required wording"], audit_rows, [Cm(5.0), Cm(2.6), Cm(8.8)], font_size=8.2)

    doc.add_heading("Appendix B｜Recommended Additional Experiments", level=2)
    exp_rows = [
        ("E1", "Modern backbone control", "VideoMAE v2 / Video Swin / SlowFast with the same LOSO split", "Shows gains are not specific to R3D-18."),
        ("E2", "History-length curve", "Last 1, 3, 5, 10, 20, all clips", "Quantifies context need and deployment memory."),
        ("E3", "Graph component ablation", "No graph / precedence only / atomic-tail / plausibility-weighted", "Isolates which graph relation helps."),
        ("E4", "External benchmark", "Comparable pre-segmented node task on HA-ViD, Assembly101, HA4M, or IMPACT", "Establishes transfer beyond thermal crimping."),
        ("E5", "Multiview/multimodal", "Single view vs three-view; RGB vs RGB+EMG/IMU", "Uses the distinctive sensing assets."),
        ("E6", "Hierarchical uncertainty", "Participant bootstrap plus within-participant run/seed resampling", "Avoids pseudo-replication and reports robust CI."),
        ("E7", "Efficiency", "Params, FLOPs, feature-cache cost, latency, memory", "Supports engineering relevance."),
        ("E8", "Annotation reliability", "Double-annotated subset with κ/α and boundary deviation", "Supports dataset validity."),
    ]
    add_table(doc, ["ID", "Experiment", "Minimum design", "Reviewer question answered"], exp_rows, [Cm(1.0), Cm(3.5), Cm(6.5), Cm(5.0)], font_size=8.0)


def main():
    pipeline = ASSETS / "method_pipeline.png"
    results = ASSETS / "main_results.png"
    make_pipeline_figure(pipeline)
    make_results_figure(results)

    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "Thermal Crimping Task-Graph Action Recognition: Submission Assessment and Draft"
    props.subject = "IEEE Access readiness assessment and English manuscript draft"
    props.author = "Draft prepared from project artifacts"
    props.keywords = "thermal crimping; action recognition; task graph; atomic tail; LOSO"
    props.comments = "Content draft; complete bracketed items before submission."

    add_title_page(doc)
    add_assessment(doc)
    add_manuscript(doc, pipeline, results)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
